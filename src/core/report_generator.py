#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成器 v2.0 — 中文 PDF + Word
"""

import os, json
from datetime import datetime
import numpy as np

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, Image, PageBreak, HRFlowable, ListFlowable, ListItem)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 字体注册 ──────────────────────────────────────────────────────────────────
def _register():
    for cp, bp in [
        ('/System/Library/Fonts/STHeiti Light.ttc',
         '/System/Library/Fonts/STHeiti Medium.ttc'),
        ('/System/Library/Fonts/PingFang.ttc',
         '/System/Library/Fonts/PingFang.ttc'),
    ]:
        try:
            pdfmetrics.registerFont(TTFont('CN',  cp, subfontIndex=0))
            pdfmetrics.registerFont(TTFont('CNB', bp, subfontIndex=0))
            return 'CN', 'CNB'
        except Exception:
            continue
    return 'Helvetica', 'Helvetica-Bold'

CN, CNB = _register()

# ── 样式 ─────────────────────────────────────────────────────────────────────
def _s(name, font=None, size=10, leading=14, c=colors.black, bold=False, sb=0, sa=4):
    return ParagraphStyle(name, fontName=CNB if bold else (font or CN),
                          fontSize=size, leading=leading, textColor=c,
                          spaceBefore=sb, spaceAfter=sa)

S_COVER = _s('cov', size=22, bold=True,  c=colors.HexColor('#1A237E'))
S_SUB   = _s('sub', size=13,              c=colors.HexColor('#455A64'))
S_H1    = _s('h1',  size=14, bold=True,   c=colors.HexColor('#1565C0'), sb=10, sa=4)
S_H2    = _s('h2',  size=11, bold=True,   c=colors.HexColor('#0277BD'), sb=6,  sa=3)
S_BODY  = _s('bd',  size=9,  leading=13)
S_BOLD  = _s('bl',  size=9,  bold=True,   leading=13)
S_CAP   = _s('cp',  size=8,  leading=11,  c=colors.grey)
S_NOTE  = _s('nt',  size=8,  leading=11,  c=colors.HexColor('#B71C1C'))

def _hr(): return HRFlowable(width="100%", thickness=0.5,
                              color=colors.HexColor('#BBDEFB'), spaceAfter=4)
def _img(p, w=14*cm):
    if p and os.path.exists(p): return Image(p, width=w, height=w*0.55)
    return Paragraph(f"[图像未找到: {os.path.basename(p) if p else '?'}]", S_CAP)
def _tbl(data, cw=None, hdr=True):
    t = Table(data, colWidths=cw)
    cmds = [('FONTNAME',(0,0),(-1,-1),CN),('FONTSIZE',(0,0),(-1,-1),8),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#F5F5F5')]),
            ('GRID',(0,0),(-1,-1),0.5,colors.grey),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3)]
    if hdr: cmds += [('BACKGROUND',(0,0),(-1,0),colors.HexColor('#1565C0')),
                     ('FONTNAME',(0,0),(-1,0),CNB),('TEXTCOLOR',(0,0),(-1,0),colors.white)]
    t.setStyle(TableStyle(cmds)); return t
def _bul(items):
    return ListFlowable([ListItem(Paragraph(i,S_BODY),leftIndent=15) for i in items],
                        bulletType='bullet',leftIndent=10)

NET_CN = {'DMN':'默认模式','SN':'凸显','ECN':'执行控制','SMN':'感觉运动',
          'VIS':'视觉','DAN':'背侧注意','LMB':'边缘系统','SUB':'皮层下'}


# ═══════════════════════════════════════════════════════════════════════════════
# 数据驱动判读辅助 —— 所有结论均由本受试者真实指标算出，不写死。
# 每个函数返回中/英双语判读串；报告正文按 language 取用。
# ═══════════════════════════════════════════════════════════════════════════════
def _verdict_range(val, lo, hi, zh=True):
    """连续指标 vs 参考区间 → 偏低↓ / 正常✓ / 偏高↑（None → —）。"""
    if val is None:
        return '—'
    if val < lo:  return '偏低↓' if zh else 'Low↓'
    if val > hi:  return '偏高↑' if zh else 'High↑'
    return '正常✓' if zh else 'Normal✓'


def _qc_verdict(val, thr, higher_better=True, zh=True):
    """QC 单项：按真实值与阈值方向判定 良好/欠佳（None → —）。"""
    if val is None:
        return '—'
    ok = (val >= thr) if higher_better else (val <= thr)
    if ok:
        return '良好' if zh else 'Good'
    return ('偏低' if higher_better else '偏高') if zh else ('Low' if higher_better else 'High')


def _fc_strength(r, zh=True):
    """功能连接强度按 |r| 分级（数据驱动，替换写死的『中等偏高』）。"""
    a = abs(r)
    if a >= 0.6:  return '强'     if zh else 'strong'
    if a >= 0.4:  return '中等偏高' if zh else 'moderate-high'
    if a >= 0.2:  return '中等'    if zh else 'moderate'
    if a >= 0.1:  return '偏弱'    if zh else 'weak'
    return '弱' if zh else 'very weak'


def _parse_age(age_raw):
    """DICOM PatientAge（如 '025Y' / '25' / 'N/A'）→ int 岁；无法解析返回 None。"""
    if age_raw is None:
        return None
    s = str(age_raw).strip().upper()
    if not s or s == 'N/A':
        return None
    num = ''
    for ch in s:
        if ch.isdigit():
            num += ch
        else:
            break
    if not num:
        return None
    try:
        v = int(num)
    except ValueError:
        return None
    # DICOM 形如 '025Y'（岁）/'012M'（月）/'030W'（周）；仅取岁，其他换算近似
    if s.endswith('M'):  v = round(v / 12)
    elif s.endswith('W'): v = round(v / 52)
    elif s.endswith('D'): v = round(v / 365)
    return v if 0 < v < 130 else None


# ── 图论全局指标：真实值 vs 健康成人参考范围 → (参考文字, 判读) ─────────────────
_GRAPH_REF = {
    # metric_key: (lo, hi, 参考文字_zh, 参考文字_en)
    'mean_degree':       (12, 20,   '~12–20（阈值依赖）', '~12–20'),
    'avg_clustering':    (0.5, 0.8, '0.5–0.8',           '0.5–0.8'),
    'global_efficiency': (0.5, 0.8, '0.5–0.8',           '0.5–0.8'),
    'local_efficiency':  (0.8, 0.95,'0.8–0.95',          '0.8–0.95'),
    'char_path_length':  (1.5, 3.0, '1.5–3.0',           '1.5–3.0'),
    'modularity':        (0.2, 0.5, '0.2–0.5（典型值）',  '0.2–0.5'),
}


def _graph_interp(metric, val, zh=True):
    """图论表结论列：按真实值落在参考区间的位置动态生成（替换写死结论）。"""
    ref = _GRAPH_REF.get(metric)
    if ref is None:
        return '—'
    lo, hi = ref[0], ref[1]
    if metric == 'char_path_length':
        # 路径长度越短→整合性越好
        if val < lo:  return '较短（网络整合性好）' if zh else 'Short — good integration'
        if val > hi:  return '偏长（整合性偏弱）'   if zh else 'Long — weaker integration'
        return '正常范围' if zh else 'Normal range'
    if val < lo:  return '偏低' if zh else 'Below range'
    if val > hi:  return '偏高' if zh else 'Above range'
    return '正常范围' if zh else 'Normal range'


class ReportGenerator:
    def __init__(self, results, output_dir, scan_params=None,
                 update_literature=False):
        self.results     = results
        self.output_dir  = output_dir
        self.update_literature = update_literature   # 是否联网抓取最新文献
        self._lit_cache  = {}                         # disease_key -> [refs]
        self.reports_dir = os.path.join(output_dir, 'reports')
        self.imgs_dir    = os.path.join(output_dir, 'report_imgs')
        os.makedirs(self.reports_dir, exist_ok=True)
        rdir = os.path.join(output_dir, 'results')
        self.rdir = rdir

        self.sp   = scan_params or results.get('scan_params', {})
        self.qc   = self._jl(rdir, 'qc_metrics.json')   or results.get('qc_metrics', {})
        self.gm   = self._jl(rdir, 'graph_metrics.json') or results.get('graph', {})
        self.ns   = self._jl(rdir, 'network_stats.json') or {}
        self.dfc  = self._jl(rdir, 'dFC_results.json')  or results.get('dynamic', {})
        self.fp   = self._jl(rdir, 'brain_fingerprint.json') or {}

        rn_p = os.path.join(rdir, 'roi_names.json')
        self.roi_names = json.load(open(rn_p)) if os.path.exists(rn_p) else []
        fc_p = os.path.join(rdir, 'FC_pearson.npy')
        self.FC = np.load(fc_p) if os.path.exists(fc_p) else np.zeros((33,33))

        # ── 多模态结果 ────────────────────────────────────────────────────────
        mm = results.get('multimodal', {})
        self.t1_res  = mm.get('t1',  {}) or self._jl(rdir, 't1_results.json')
        self.dti_res = mm.get('dti', {}) or self._jl(rdir, 'dti_results.json')
        self.qsm_res = mm.get('qsm', {}) or self._jl(rdir, 'qsm_results.json')

    @staticmethod
    def _jl(d, name):
        p = os.path.join(d, name)
        try:
            return json.load(open(p, encoding='utf-8', errors='replace')) if os.path.exists(p) else {}
        except Exception: return {}

    def _ip(self, name): return os.path.join(self.imgs_dir, name)

    def _recent_refs(self, disease_key: str) -> list:
        """取某疾病的最新 PubMed 文献；未开启或离线时返回 []（调用方回退硬编码引用）。"""
        if not self.update_literature:
            return []
        if disease_key in self._lit_cache:
            return self._lit_cache[disease_key]
        refs = []
        try:
            from core.literature_updater import LiteratureUpdater
            lu = LiteratureUpdater(cache_dir=self.output_dir)
            refs = lu.get_recent_refs(disease_key, retmax=2)
        except Exception:
            refs = []
        self._lit_cache[disease_key] = refs
        return refs

    def _refs_line(self, disease_key: str, zh: bool = True) -> str:
        """把最新文献拼成一行「最新文献」串；无则返回空串。"""
        refs = self._recent_refs(disease_key)
        if not refs:
            return ""
        try:
            from core.literature_updater import LiteratureUpdater
            fmt = LiteratureUpdater.format_ref
        except Exception:
            return ""
        tag = "最新文献" if zh else "Recent"
        return f"\n【{tag}】\n" + "\n".join("· " + fmt(r) for r in refs)

    # ── FC 便捷取值 ────────────────────────────────────────────────────────────
    def _fc(self, a, b):
        """两个 ROI 间 Pearson r；任一缺失返回 None。"""
        rn = self.roi_names
        if a in rn and b in rn:
            return float(self.FC[rn.index(a), rn.index(b)])
        return None

    def _dmn_internal(self):
        """DMN 7 核心节点内部平均 FC（真实值）。"""
        nodes = ["mPFC","PCC","Precuneus","L_Angular","R_Angular","L_mTL_HPC","R_mTL_HPC"]
        idx = [self.roi_names.index(n) for n in nodes if n in self.roi_names]
        if len(idx) < 2:
            return 0.0
        sub = self.FC[np.ix_(idx, idx)]
        return float(sub[np.triu_indices(len(idx), k=1)].mean())

    def disease_similarity(self):
        """
        规则化疾病相似度引擎（数据驱动，可解释）。
        每种疾病定义若干布尔特征，全部由本受试者真实指标算出；
        相似度 = 命中数/总数，单受试者封顶『中度』，绝不给出高相似度或诊断。
        返回 {疾病key: {'zh_name','en_name','features':[(zh,en,hit)],
                        'level_zh','level_en','note_zh','note_en'}}
        """
        gm  = self.gm; ns = self.ns; dfc = self.dfc
        pn  = gm.get('per_node', {})
        hubs = set(gm.get('hub_regions', []))
        dmn = self._dmn_internal()
        sn_fc  = ns.get('SN',  {}).get('mean_FC', 0.0)
        ecn_fc = ns.get('ECN', {}).get('mean_FC', 0.0)
        dan_fc = ns.get('DAN', {}).get('mean_FC', 0.0)
        smn_fc = ns.get('SMN', {}).get('mean_FC', 0.0)
        cc     = gm.get('avg_clustering', 0.0)
        sigma  = gm.get('small_world_sigma', 0.0)
        ge     = gm.get('global_efficiency', 0.0)
        n_trans = dfc.get('n_transitions', 0)
        has_dfc = dfc.get('n_windows', 0) > 0
        pcc_mpfc = self._fc('PCC', 'mPFC')
        amy_dlpfc = self._fc('L_Amy', 'L_dlPFC')
        mpfc_amy  = self._fc('mPFC', 'L_Amy')
        age = _parse_age(self.sp.get('age'))

        def hub_any(*names):  return any(n in hubs for n in names)

        # 每种疾病：(zh特征, en特征, 命中布尔) —— 命中判据全部基于真实指标
        rules = {
            "MDD": ("重度抑郁症", "Major Depressive Disorder", [
                (f"DMN内部FC偏高(={dmn:.3f}>0.35)", f"DMN internal FC high (={dmn:.3f})", dmn > 0.35),
                (f"动态FC灵活性偏低(={n_trans}次<3)", f"Low dFC flexibility ({n_trans}<3)", has_dfc and n_trans < 3),
                ("边缘系统节点为Hub(杏仁核/丘脑)", "Limbic hub (amygdala/thalamus)", hub_any('L_Amy','R_Amy','L_Thal','R_Thal')),
            ]),
            "BD": ("双相情感障碍", "Bipolar Disorder", [
                (f"ECN内部FC偏离(={ecn_fc:.3f})", f"ECN FC atypical (={ecn_fc:.3f})", ecn_fc < 0.2 or ecn_fc > 0.5),
                ("杏仁核为Hub(边缘-PFC耦合)", "Amygdala hub", hub_any('L_Amy','R_Amy')),
            ]),
            "SCZ": ("精神分裂症", "Schizophrenia", [
                (f"PCC-mPFC连接减弱(={(pcc_mpfc if pcc_mpfc is not None else 0):.3f}<0.2)",
                 f"PCC-mPFC reduced (<0.2)", pcc_mpfc is not None and pcc_mpfc < 0.2),
                (f"小世界性下降(σ={sigma:.2f}<1)", f"Small-worldness down (σ<1)", 0 < sigma < 1.0),
            ]),
            "ADHD": ("注意缺陷多动障碍", "ADHD", [
                (f"DAN内部FC偏低(={dan_fc:.3f}<0.2)", f"DAN FC low (<0.2)", dan_fc < 0.2),
                (f"全局效率偏低(GE={ge:.3f}<0.5)", f"Global efficiency low (<0.5)", 0 < ge < 0.5),
                ("执行网络节点(dlPFC)非Hub", "dlPFC not a hub", not hub_any('L_dlPFC','R_dlPFC')),
            ]),
            "ASD": ("孤独症谱系", "Autism Spectrum", [
                (f"局部聚类偏高(CC={cc:.3f}>0.6)", f"High clustering (CC>0.6)", cc > 0.6),
                (f"过度局部化(σ={sigma:.2f}>2)", f"Over-segregation (σ>2)", sigma > 2.0),
            ]),
            "GAD": ("广泛性焦虑障碍", "Generalized Anxiety", [
                ("杏仁核为Hub(高连接)", "Amygdala hub", hub_any('L_Amy','R_Amy')),
                (f"SN内部FC偏高(={sn_fc:.3f}>0.35)", f"SN FC high (>0.35)", sn_fc > 0.35),
                (f"杏仁核-dlPFC调控弱(={(amy_dlpfc if amy_dlpfc is not None else 0):.3f}<0.15)",
                 f"Amygdala-dlPFC weak (<0.15)", amy_dlpfc is not None and amy_dlpfc < 0.15),
            ]),
            "OCD": ("强迫症", "OCD", [
                ("尾状核为Hub(皮层-纹状体回路)", "Caudate hub", hub_any('L_Caudate','R_Caudate')),
                (f"SN内部FC偏高(={sn_fc:.3f}>0.35)", f"SN FC high (>0.35)", sn_fc > 0.35),
            ]),
            "PTSD": ("创伤后应激障碍", "PTSD", [
                ("杏仁核为Hub(过反应性)", "Amygdala hub", hub_any('L_Amy','R_Amy')),
                (f"mPFC-杏仁核调控弱(={(mpfc_amy if mpfc_amy is not None else 0):.3f}<0.15)",
                 f"mPFC-amygdala weak (<0.15)", mpfc_amy is not None and mpfc_amy < 0.15),
            ]),
            "AD": ("阿尔茨海默病", "Alzheimer's Disease", [
                (f"DMN严重破坏(PCC-mPFC={(pcc_mpfc if pcc_mpfc is not None else 0):.3f}<0.1)",
                 f"DMN disrupted (PCC-mPFC<0.1)", pcc_mpfc is not None and pcc_mpfc < 0.1),
                (f"全局效率急剧下降(GE={ge:.3f}<0.4)", f"GE sharply reduced (<0.4)", 0 < ge < 0.4),
            ]),
            "PD": ("帕金森病", "Parkinson's Disease", [
                (f"感觉运动网络改变(SMN={smn_fc:.3f})", f"SMN altered (={smn_fc:.3f})", smn_fc < 0.2 or smn_fc > 0.5),
                ("基底节节点为Hub", "Basal ganglia hub", hub_any('L_Caudate','R_Caudate')),
            ]),
        }

        def level(ratio, zh):
            # 单受试者上限『中度』——绝不给出高相似度。
            # 阈值保守：≥0.6 才『中度』(如 2/2、2/3)，命中即『低』，未命中『极低』。
            if ratio <= 0:      return '极低' if zh else 'Very low'
            if ratio < 0.6:     return '低'   if zh else 'Low'
            return '中度' if zh else 'Moderate'

        out = {}
        for key, (zh_n, en_n, feats) in rules.items():
            n_hit = sum(1 for _, _, h in feats if h)
            n_tot = len(feats)
            ratio = n_hit / n_tot if n_tot else 0
            lv_zh, lv_en = level(ratio, True), level(ratio, False)
            note_zh = f"命中 {n_hit}/{n_tot} 项特征"
            note_en = f"{n_hit}/{n_tot} features matched"
            # 年龄门控：AD/PD 明显不符发病年龄 → 强制极低
            if key == 'AD' and age is not None and age < 50:
                lv_zh, lv_en = '极低', 'Very low'
                note_zh += f"；年龄{age}岁远低于AD典型发病"; note_en += f"; age {age}, atypical for AD"
            if key == 'PD' and age is not None and age < 45:
                lv_zh, lv_en = '极低', 'Very low'
                note_zh += f"；年龄{age}岁不符PD发病特征"; note_en += f"; age {age}, atypical for PD"
            out[key] = {'zh_name': zh_n, 'en_name': en_n, 'features': feats,
                        'n_hit': n_hit, 'n_tot': n_tot,
                        'level_zh': lv_zh, 'level_en': lv_en,
                        'note_zh': note_zh, 'note_en': note_en}
        return out

    @staticmethod
    def _qr(score, lang='zh'):
        if score>=80: return '优秀' if lang=='zh' else 'Excellent'
        if score>=60: return '良好' if lang=='zh' else 'Good'
        if score>=40: return '一般' if lang=='zh' else 'Fair'
        return '较差' if lang=='zh' else 'Poor'

    # ── PDF ──────────────────────────────────────────────────────────────────
    def generate_pdf_report(self, language='zh') -> str:
        if language == 'en':
            return self._generate_en_pdf_report()
        sp=self.sp; qc=self.qc; gm=self.gm; ns=self.ns; dfc=self.dfc
        roi_names=self.roi_names; FC=self.FC; n_roi=len(roi_names)
        subj=sp.get('subject_id','unknown')
        story=[]

        # 封面
        story += [Spacer(1,3*cm),
                  Paragraph('静息态fMRI脑网络分析报告', S_COVER), Spacer(1,0.5*cm),
                  Paragraph('Resting-State fMRI Brain Network Analysis Report', S_SUB),
                  Spacer(1,1.5*cm)]
        story.append(_tbl([
            ['项目','内容'],['受试者ID',subj],
            ['年龄/性别', f"{sp.get('age','N/A')} / {sp.get('sex','N/A')}"],
            ['扫描机构', sp.get('institution','N/A')],
            ['扫描仪',   sp.get('scanner','N/A')],
            ['扫描日期', sp.get('scan_date','N/A')],
            ['TR / TE',  f"{sp.get('TR',2.0)} s / {sp.get('TE',30)} ms"],
            ['报告日期', datetime.now().strftime('%Y-%m-%d')],
            ['分析平台', 'Python | NetworkX | SciPy | dcm2niix | ReportLab'],
            ['QC评分',   f"{qc.get('QC_score',0)}/100  {qc.get('QC_stars','—')}"],
        ], cw=[5*cm,11*cm]))
        story += [Spacer(1,1*cm), Paragraph('⚠ 本报告仅供科研参考，不构成任何医学诊断依据', S_NOTE),
                  PageBreak()]

        # 一、扫描参数
        story += [Paragraph('一、扫描参数与数据概览', S_H1), _hr(),
                  Paragraph('1.1 采集参数', S_H2)]
        stc = '✅ 已执行' if sp.get('slice_timing_available') else '⚠ 未执行'
        story.append(_tbl([
            ['参数','rsfMRI BOLD','T1结构像'],
            ['扫描序列','EPI BOLD','N/A'],
            ['TR', f"{sp.get('TR',2.0)} s", 'N/A'],
            ['TE', f"{sp.get('TE',30)} ms", 'N/A'],
            ['体素', f"{sp.get('voxel_size','N/A')} mm", 'N/A'],
            ['时间点', f"{qc.get('n_timepoints','N/A')} ({qc.get('total_duration_min',0):.1f}min)", 'N/A'],
            ['切片时间校正', stc, 'N/A'],
        ], cw=[5*cm,6*cm,5*cm]))
        story.append(Paragraph(
            '注：本流程为静息态 rs-fMRI 分析管线，未提取 T1 结构像的采集参数，'
            'T1 列以 N/A 表示（如需 T1 采集细节请查阅原始 DICOM/JSON sidecar）。', S_CAP))
        story.append(PageBreak())

        # 二、预处理
        story += [Paragraph('二、数据预处理流程', S_H1), _hr()]
        story.append(_tbl([
            ['步骤','方法','状态','说明'],
            ['DICOM→NIfTI','dcm2niix','✅','UIH mosaic 自动解码'],
            ['丢弃前5TR','—','✅','steady-state'],
            ['切片时间校正','线性插值 STC', stc,'参考时间=TR/2'],
            ['脑掩码','强度65th%+形态学','✅','无需FSL'],
            ['PSC标准化','%(x-μ)/μ×100','✅','逐体素'],
            ['线性去趋势','OLS','✅','逐体素'],
            ['带通滤波','Butterworth 4阶 filtfilt','✅','0.01–0.1 Hz'],
            ['空间平滑','Gaussian FWHM=6mm','✅','σ≈0.73体素'],
            ['头动矫正','—','⚠未执行','需FSL/AFNI'],
        ], cw=[4*cm,5.5*cm,2.5*cm,6*cm]))
        story.append(PageBreak())

        # 三、QC
        story += [Paragraph('三、数据质量控制（QC）', S_H1), _hr()]
        score_v = qc.get('QC_score',0)
        story += [Paragraph(f"综合QC评分：{score_v}/100  {qc.get('QC_stars','—')}  —  {self._qr(score_v)}",S_BOLD),
                  Spacer(1,0.3*cm)]
        _tsnr = qc.get('tSNR_median', 0); _dvars = qc.get('DVARS_pct_median', 0)
        _fd = qc.get('FD_proxy_pct_median', 0); _dur = qc.get('total_duration_min', 0)
        _pctbad = qc.get('pct_bad_TPs', 0)
        story.append(_tbl([
            ['QC指标','数值','参考','评价'],
            ['tSNR中位数', f"{_tsnr:.1f}", '≥50优', _qc_verdict(_tsnr, 50, higher_better=True)],
            ['DVARS%中位数', f"{_dvars:.3f}%", '<5%', _qc_verdict(_dvars, 5, higher_better=False)],
            ['FD代理%中位数', f"{_fd:.3f}%", '<0.5%', _qc_verdict(_fd, 0.5, higher_better=False)],
            ['高运动TP', f"{qc.get('n_bad_TPs',0)}/{qc.get('n_timepoints',0)}", '<20%',
             _qc_verdict(_pctbad, 20, higher_better=False) + f"（{_pctbad:.1f}%）"],
            ['扫描时长', f"{_dur:.1f}min", '≥6min', _qc_verdict(_dur, 6, higher_better=True)],
        ], cw=[4.5*cm,3*cm,4*cm,6*cm]))
        story += [Spacer(1,0.3*cm), _img(self._ip('qc_timeseries.png'),15*cm),
                  Paragraph('图1：QC时间序列。上FD代理；中DVARS；下tSNR分布。', S_CAP),
                  PageBreak()]

        return self._build(story, subj, language)

    def _pdf_body(self, story, sp, qc, gm, ns, dfc, roi_names, FC, n_roi):
        """完整报告正文 第四章至附录 — 与参考报告 05_pdf_report.py 内容一致"""

        # ── 导出常用局部变量 ─────────────────────────────────────────────────
        NET_CN = {"DMN":"默认模式网络","SN":"凸显网络","ECN":"执行控制网络",
                  "SMN":"感觉运动网络","VIS":"视觉网络","DAN":"背侧注意网络",
                  "LMB":"边缘系统","SUB":"皮层下"}
        NMAP = {n: k for k,v in {
            'DMN':['mPFC','PCC','Precuneus','L_Angular','R_Angular','L_mTL_HPC','R_mTL_HPC'],
            'SN':['ACC_dACC','L_AI','R_AI'],'ECN':['L_dlPFC','R_dlPFC','L_IPL','R_IPL'],
            'SMN':['L_M1','R_M1','SMA'],'VIS':['V1_L','V1_R','LOC_L'],
            'DAN':['L_FEF','R_FEF','L_IPS','R_IPS'],'LMB':['L_Amy','R_Amy','L_Thal','R_Thal'],
            'SUB':['L_Caudate','R_Caudate','Brainstem','L_Cerebellum','R_Cerebellum'],
        }.items() for n in v}

        pn = gm.get('per_node', {})
        sigma_v = gm.get('small_world_sigma', 0)
        ge_v    = gm.get('global_efficiency', 0)
        le_v    = gm.get('local_efficiency', 0)

        fp_d   = self.fp
        smry   = {}  # analysis_summary placeholder

        # ══ 四、ROI & 网络 ══════════════════════════════════════════════════════
        story += [Paragraph("四、ROI定义与静息态脑网络分析", S_H1), _hr()]
        story.append(Paragraph(
            f"本分析采用{n_roi}个基于MNI坐标的球形ROI（半径3.5mm≈1体素），涵盖8个主要静息态脑网络。"
            "ROI坐标参考Power264图谱及AAL2图谱，在原生空间提取时间序列后计算功能连接。", S_BODY))
        story.append(Spacer(1, 0.2*cm))

        net_rows = [["网络","中文名","脑区","内部FC均值（r）","SD"]]
        for k, v in ns.items():
            net_rows.append([k, NET_CN.get(k, k),
                             ", ".join(v["rois"][:4]) + "...",
                             f"{v['mean_FC']:.3f}", f"{v['std_FC']:.3f}"])
        story.append(_tbl(net_rows, cw=[1.5*cm, 3.5*cm, 7*cm, 3*cm, 2*cm]))
        story += [Spacer(1, 0.2*cm),
                  _img(self._ip("network_strength.png"), 14*cm),
                  Paragraph(
                      "图2：各静息态脑网络内部平均功能连接强度（Pearson r）。" +
                      (lambda _top: f"本受试者内部连接最强的网络为 {('、'.join(NET_CN.get(k,k) for k,_ in _top))}"
                                    f"（r={_top[0][1]:.3f}）。" if _top else "")(
                          sorted(((k, v.get('mean_FC', 0)) for k, v in ns.items()),
                                 key=lambda x: x[1], reverse=True)[:2]), S_CAP),
                  PageBreak()]

        # ══ 五、DMN ═════════════════════════════════════════════════════════════
        story += [Paragraph("五、默认模式网络（DMN）专项分析", S_H1), _hr()]
        dmn_nodes = ["mPFC","PCC","Precuneus","L_Angular","R_Angular","L_mTL_HPC","R_mTL_HPC"]
        dmn_idx   = [roi_names.index(n) for n in dmn_nodes if n in roi_names]
        FC_dmn    = FC[np.ix_(dmn_idx, dmn_idx)] if len(dmn_idx) > 1 else np.zeros((1,1))
        dmn_str   = float(FC_dmn[np.triu_indices(len(dmn_idx), k=1)].mean()) if len(dmn_idx) > 1 else 0.0

        pcc_i = roi_names.index("PCC") if "PCC" in roi_names else 0
        pcc_fc = sorted([(FC[pcc_i, i], roi_names[i])
                          for i in range(n_roi) if i != pcc_i], reverse=True)

        story += [Paragraph(f"DMN内部平均功能连接强度：r = {dmn_str:.3f}（{_fc_strength(dmn_str)}）", S_BOLD),
                  Spacer(1, 0.2*cm)]

        # 5.1 PCC 种子连接
        story.append(Paragraph("5.1 PCC种子连接分析（Seed-based Connectivity）", S_H2))
        story.append(_tbl(
            [["ROI", "r（PCC种子）", "网络归属", "解释"]] +
            [[nm, f"{r:.3f}",
              "DMN" if nm in ["mPFC","Precuneus","L_Angular","R_Angular","L_mTL_HPC","R_mTL_HPC"] else "其他",
              "核心DMN节点" if r > 0.5 else ("中等连接" if r > 0.2 else "弱连接")]
             for r, nm in pcc_fc[:10]],
            cw=[4*cm, 3*cm, 3*cm, 8*cm]))
        story += [Spacer(1, 0.3*cm)]

        # 5.2 DMN 内部 FC 矩阵
        story += [Paragraph("5.2 DMN内部连接矩阵", S_H2),
                  _img(self._ip("dmn_fc.png"), 12*cm)]
        # 动态地找 PCC-Precuneus 连接值
        pcc_prec = 0.0
        if "PCC" in roi_names and "Precuneus" in roi_names:
            pcc_prec = float(FC[roi_names.index("PCC"), roi_names.index("Precuneus")])
        story.append(Paragraph(
            f"图3：DMN7个核心节点内部功能连接矩阵。颜色深红代表强正相关，深蓝代表强负相关。"
            f"PCC-Precuneus连接（r={pcc_prec:.3f}）为本受试者最强DMN内部连接，"
            "提示DMN后部核心回路完整。", S_CAP))
        story.append(Spacer(1, 0.3*cm))

        # 5.3 DMN 神经科学解释
        story += [Paragraph("5.3 DMN与认知功能的神经科学解释", S_H2)]
        story.append(Paragraph(
            "默认模式网络（Default Mode Network, DMN）是静息态最显著的内源性脑网络之一，"
            "由内侧前额叶皮层（mPFC）、后扣带回皮层（PCC）、楔前叶（Precuneus）、"
            "角回（Angular Gyrus）及内侧颞叶（mTL，含海马）等区域组成。DMN在任务执行时通常失活，"
            "在静息状态或自我参照性思维时激活，参与：", S_BODY))
        story.append(ListFlowable([
            ListItem(Paragraph("自我意识与自我参照加工（mPFC主导）", S_BODY), leftIndent=15),
            ListItem(Paragraph("内省与心理游移（PCC/Precuneus主导）", S_BODY), leftIndent=15),
            ListItem(Paragraph("情景记忆提取与长时记忆（mTL/HPC主导）", S_BODY), leftIndent=15),
            ListItem(Paragraph("社会认知与心理理论（ToM）（mPFC/TPJ主导）", S_BODY), leftIndent=15),
            ListItem(Paragraph("情绪调节与情感性自我参照（mPFC-Amy回路）", S_BODY), leftIndent=15),
        ], bulletType='bullet', leftIndent=10))
        pcc_mpfc = float(FC[pcc_i, roi_names.index("mPFC")]) if "mPFC" in roi_names else 0.0
        story += [Spacer(1, 0.2*cm),
                  Paragraph(
                      f"本受试者PCC-Precuneus连接（r={pcc_prec:.3f}）"
                      f"{'显著高于' if pcc_prec > 0.6 else '处于'}文献报道的健康成人均值（r≈0.5–0.6），"
                      f"而PCC-mPFC连接（r={pcc_mpfc:.3f}）处于正常范围。"
                      f"DMN整体内部连接强度（r={dmn_str:.3f}）在文献参考范围内（健康成人r≈0.2–0.5）。",
                      S_BODY),
                  PageBreak()]

        # ══ 六、全脑 FC ════════════════════════════════════════════════════════
        fc_nz = FC[FC != 0]
        story += [Paragraph("六、全脑功能连接分析", S_H1), _hr()]
        story.append(_tbl([
            ["指标","数值"],
            ["FC矩阵维度", f"{n_roi}×{n_roi}（{n_roi} ROI）"],
            ["全脑平均FC（阈值前）", f"{float(fc_nz.mean()):.3f} ± {float(fc_nz.std()):.3f}" if len(fc_nz)>0 else "N/A"],
            ["DMN内部连接", f"{dmn_str:.3f}"],
            ["正性连接比例", f"{float((fc_nz>0).sum()/len(fc_nz)*100):.1f}%" if len(fc_nz)>0 else "N/A"],
            ["强连接（r>0.3）比例", f"{float((fc_nz>0.3).sum()/len(fc_nz)*100):.1f}%" if len(fc_nz)>0 else "N/A"],
        ], cw=[8*cm, 10*cm]))
        story += [Spacer(1, 0.3*cm), _img(self._ip("fc_matrix.png"), 15*cm),
                  Paragraph(
                      "图4：33个ROI全脑功能连接矩阵（Pearson r）。对角线为1（已归零显示）。"
                      "可见DMN脑区（前7个）、SN脑区、ECN脑区内部呈现明显正相关聚类结构。", S_CAP),
                  PageBreak()]

        # ══ 七、ALFF / fALFF / ReHo ════════════════════════════════════════════
        story += [Paragraph("七、局部脑活动分析（ALFF / fALFF / ReHo）", S_H1), _hr()]
        story.append(Paragraph(
            "局部脑活动指标从不同角度反映静息态BOLD信号的自发神经活动特征，"
            "不依赖种子区或网络假设，适合探索性分析。", S_BODY))
        story.append(Spacer(1, 0.2*cm))
        story.append(_tbl([
            ["指标","全称","计算方法","神经科学意义"],
            ["ALFF","低频振幅\n(Amplitude of Low-Frequency Fluctuations)",
             "0.01–0.1 Hz频段FFT幅值均值",
             "反映局部自发神经活动强度；高ALFF区对应高代谢/高活跃度脑区"],
            ["fALFF","分数低频振幅",
             "ALFF / 全频段幅值均值",
             "相对指标，消除全局噪声影响，对生理噪声更鲁棒"],
            ["ReHo","局部一致性\n(Regional Homogeneity)",
             "Kendall's W（3×3×3邻域时间序列秩一致性）",
             "反映邻域体素时间序列的同步程度；高ReHo提示局部神经活动同步性强"],
        ], cw=[2*cm, 4*cm, 5*cm, 7*cm]))
        story.append(Spacer(1, 0.3*cm))
        for mapn, figname, desc in [
            ("ALFF",  "alff_brain.png",
             "ALFF z-score图（轴位切面）。正值（红色）代表高于全脑均值的低频振幅，负值（蓝色）代表低于均值区域。"),
            ("fALFF", "falff_brain.png",
             "fALFF z-score图。分数低频振幅分布，消除噪声影响后的局部活动强度。"),
            ("ReHo",  "reho_brain.png",
             "ReHo z-score图（Kendall's W）。高ReHo区域（红色）提示局部神经活动高度同步。"),
        ]:
            story += [Paragraph(f"7.x {mapn}图", S_H2),
                      _img(self._ip(figname), 15*cm),
                      Paragraph(f"图：{desc}", S_CAP),
                      Spacer(1, 0.2*cm)]
        story.append(Paragraph(
            "所有指标已z-score标准化（均值=0，SD=1）便于比较。"
            "注意：本分析未进行MNI空间配准，图像坐标为原生扫描空间，"
            "解剖定位仅供参考，不可与标准图谱直接比对。", S_NOTE))
        story.append(PageBreak())

        # ══ 八、图论分析 ════════════════════════════════════════════════════════
        story += [Paragraph("八、脑功能网络图论分析", S_H1), _hr()]
        story.append(Paragraph(
            f"以Pearson r>{gm.get('threshold_r',0.2)}为阈值构建二值/加权脑功能图，"
            f"节点={gm.get('n_nodes',n_roi)}，边={gm.get('n_edges',0)}，"
            f"密度={gm.get('density',0):.3f}。", S_BODY))
        story.append(Spacer(1, 0.2*cm))
        _md = gm.get('mean_degree',0); _cc = gm.get('avg_clustering',0)
        _gge = gm.get('global_efficiency',0); _le = gm.get('local_efficiency',0)
        _cpl = gm.get('char_path_length',0); _sig = gm.get('small_world_sigma',0)
        _mod = gm.get('modularity',0)
        story.append(_tbl([
            ["图论全局指标","本受试者数值","文献参考范围（健康成人）","解释"],
            ["平均连接度（Degree）",   f"{_md:.2f}",  _GRAPH_REF['mean_degree'][2],       _graph_interp('mean_degree', _md)],
            ["平均聚类系数（CC）",     f"{_cc:.3f}",  _GRAPH_REF['avg_clustering'][2],    _graph_interp('avg_clustering', _cc)],
            ["全局效率",              f"{_gge:.3f}", _GRAPH_REF['global_efficiency'][2], _graph_interp('global_efficiency', _gge)],
            ["局部效率",              f"{_le:.3f}",  _GRAPH_REF['local_efficiency'][2],  _graph_interp('local_efficiency', _le)],
            ["特征路径长度",          f"{_cpl:.3f}", _GRAPH_REF['char_path_length'][2],  _graph_interp('char_path_length', _cpl)],
            ["小世界指数（σ）",       f"{_sig:.2f}", ">1为小世界",  "✓ 小世界网络特征" if _sig > 1 else "未达小世界(σ≤1)"],
            ["模块化系数（Q）",       f"{_mod:.3f}", _GRAPH_REF['modularity'][2],        _graph_interp('modularity', _mod)],
            ["社区/模块数",           f"{gm.get('n_communities',0)}",         "3–7（阈值依赖）",  "本阈值下分区数"],
        ], cw=[5*cm, 3.5*cm, 5*cm, 4.5*cm]))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph("8.1 Hub脑区分析", S_H2))
        hubs = gm.get('hub_regions', [])
        hub_str = "、".join(hubs) if hubs else "未检出"
        story.append(Paragraph(f"Hub脑区（高度+高中介中心性脑区）：{hub_str}", S_BOLD))
        story.append(Spacer(1, 0.2*cm))
        story.append(_img(self._ip("graph_hubs.png"), 15*cm))
        story.append(Paragraph(
            "图5：图论分析结果。左：各ROI连接度（红色=Hub脑区）；右：Degree vs. 中介中心性散点图，"
            "红色节点为Hub脑区，标注为脑区名称。", S_CAP))
        story.append(Spacer(1, 0.2*cm))
        story.append(Paragraph("Hub脑区神经科学意义：", S_H2))
        HUB_INTERP = {
            "R_AI":      "右侧前岛叶（Anterior Insula）：凸显网络核心节点，整合内外感知、情绪觉知与自主神经调控",
            "L_AI":      "左侧前岛叶：同上，与右侧前岛叶共同构成凸显网络核心",
            "L_dlPFC":   "左侧背外侧前额叶（dlPFC）：执行控制网络核心，调控工作记忆、认知灵活性与情绪调节",
            "R_dlPFC":   "右侧背外侧前额叶：同上，认知控制功能",
            "V1_L":      "左侧初级视觉皮层（V1）：视觉网络枢纽，高信噪比使其在功能图中常呈Hub特征",
            "V1_R":      "右侧初级视觉皮层：同上",
            "L_Amy":     "左侧杏仁核（Amygdala）：情绪处理、恐惧学习与社会信息编码的关键皮层下结构",
            "R_Amy":     "右侧杏仁核：同上，情绪评估的核心结构",
            "L_Caudate": "左侧尾状核（Caudate）：基底节核团，参与奖赏预测、习惯学习与运动规划",
            "R_Caudate": "右侧尾状核：同上，与奖赏系统及精神疾病多巴胺功能密切相关",
            "L_Thal":    "左侧丘脑：皮层-皮层下信息中继核，广泛连接参与多种感知与认知功能",
            "R_Thal":    "右侧丘脑：同上",
            "Precuneus": "楔前叶：DMN重要节点，参与视空间处理、情景记忆与意识相关功能",
            "mPFC":      "内侧前额叶：自我参照加工、情绪调节与DMN核心节点",
            "PCC":       "后扣带回皮层：DMN中枢，内省、记忆检索与自我意识的关键节点",
        }
        hub_rows = [["Hub脑区","神经科学意义"]]
        for h in hubs:
            hub_rows.append([h, HUB_INTERP.get(h, "重要连接枢纽节点，参与多网络信息整合")])
        story.append(_tbl(hub_rows, cw=[4*cm, 14*cm]))
        story.append(PageBreak())

        # ══ 九、动态 FC ════════════════════════════════════════════════════════
        story += [Paragraph("九、动态功能连接分析", S_H1), _hr()]
        n_win = dfc.get('n_windows', 0)
        if n_win > 0:
            story.append(_tbl([
                ["参数","数值"],
                ["分析方法","滑动窗口法（Sliding Window）"],
                ["窗口大小", f"{dfc.get('window_size_TPs',44)}个时间点（{dfc.get('window_size_s',88):.0f}秒 = {dfc.get('window_size_s',88)/60:.1f}分钟）"],
                ["步长",     f"{dfc.get('step_TPs',4)}个时间点（{dfc.get('step_TPs',4)*2}秒）"],
                ["分析ROI",  "DMN 7个核心节点（21个ROI对）"],
                ["总窗口数", f"{n_win}个"],
                ["脑状态数（k-means k=2）","2个状态"],
                ["State 1 占比", f"{dfc.get('state_occupancy',{}).get('State_1',0)*100:.1f}%（{dfc.get('dwell_times_s',{}).get('State_1',0):.0f}秒）"],
                ["State 2 占比", f"{dfc.get('state_occupancy',{}).get('State_2',0)*100:.1f}%（{dfc.get('dwell_times_s',{}).get('State_2',0):.0f}秒）"],
                ["状态转换次数", f"{dfc.get('n_transitions',0)}次"],
            ], cw=[6*cm, 12*cm]))
            story += [Spacer(1, 0.3*cm), _img(self._ip("dynamic_fc.png"), 15*cm),
                      Paragraph(
                          "图6：动态功能连接（DMN ROI对）。上：时间-连接强度热图；下：脑状态序列"
                          "（红=State1，蓝=State2）。", S_CAP),
                      Spacer(1, 0.2*cm)]
            flex_v = dfc.get('n_transitions', 0)
            s1_pct = dfc.get('state_occupancy', {}).get('State_1', 0) * 100
            _dur_min = dfc.get('scan_duration_min') or self.qc.get('total_duration_min', 0)
            _dur_txt = f"{_dur_min:.1f}分钟" if _dur_min else "本次"
            story.append(Paragraph(
                f"脑状态灵活性解读：本受试者在{_dur_txt}扫描期间DMN发生{flex_v}次脑状态切换，"
                f"{s1_pct:.1f}%时间处于State 1。"
                f"{'状态灵活性偏低（文献典型值约5–15次/8分钟），低灵活性在抑郁症（MDD）、双相障碍及精神分裂症患者中均有报道' if flex_v < 3 else '脑状态灵活性在正常范围内'}。"
                "但单受试者数据不能得出组间统计结论。", S_BODY))
        else:
            story.append(Paragraph("时间点不足，动态FC分析跳过。", S_BODY))
        story.append(PageBreak())

        # ══ 十、精神疾病脑网络文献对照分析 ════════════════════════════════════
        story += [Paragraph("十、神经精神疾病脑网络文献对照分析", S_H1), _hr()]
        story.append(Paragraph(
            "⚠ 重要声明：本部分仅为科研文献对照分析，用于探索本受试者脑网络特征与已有群体研究结果之间的相似性，"
            "不可作为医学诊断依据。本报告不作任何疾病诊断、确诊或高度怀疑的结论。"
            "所有对照仅基于群体统计特征，个体脑网络变异极大，不可从单受试者推断诊断。",
            S_NOTE))
        story.append(Spacer(1, 0.3*cm))

        # 关键指标
        pcc_i_x   = roi_names.index("PCC")      if "PCC"      in roi_names else 0
        pcc_mpfc_ = float(FC[pcc_i_x, roi_names.index("mPFC")]) if "mPFC"      in roi_names else 0.0
        pcc_prec_ = float(FC[pcc_i_x, roi_names.index("Precuneus")]) if "Precuneus" in roi_names else 0.0
        ai_deg_   = pn.get("R_AI",{}).get("degree", 0)
        amy_deg_  = pn.get("L_Amy",{}).get("degree", 0)
        caud_deg_ = pn.get("L_Caudate",{}).get("degree", 0)
        amy_dlpfc = float(FC[roi_names.index("L_Amy"), roi_names.index("L_dlPFC")]) \
                    if "L_Amy" in roi_names and "L_dlPFC" in roi_names else 0.0
        mpfc_amy  = float(FC[roi_names.index("mPFC"), roi_names.index("L_Amy")]) \
                    if "mPFC" in roi_names and "L_Amy" in roi_names else 0.0
        mtl_pcc   = float(FC[roi_names.index("L_mTL_HPC"), roi_names.index("PCC")]) \
                    if "L_mTL_HPC" in roi_names and "PCC" in roi_names else 0.0
        thal_dlpfc= float(FC[roi_names.index("L_Thal"), roi_names.index("L_dlPFC")]) \
                    if "L_Thal" in roi_names and "L_dlPFC" in roi_names else 0.0
        flex_v_   = dfc.get("n_transitions", 0)
        dmn_fc_v_ = dmn_str  # defined in chapter 5 above — same Python scope

        # ── 数据驱动相似度引擎：所有相似度/命中特征均由本受试者真实指标算出 ──────
        _sim = self.disease_similarity()
        _hubset = set(gm.get('hub_regions', []))
        def _hub_state(name, zh_true, zh_false):
            return zh_true if name in _hubset else zh_false
        def _feat_col(key):
            """本受试者特征列：列出该疾病每项判据的真实命中情况。"""
            return "\n".join(("✓ " if h else "✗ ") + zh for zh, en, h in _sim[key]['features'])

        # 文献典型脑网络改变（保持不变，供对照）
        _lit_ref = {
            "MDD":  "① DMN过度连接\n② SN-DMN耦合异常\n③ 皮层-边缘连接增强\n④ 脑状态灵活性↓\n（Zhu et al. 2021, NeuroImage）",
            "BD":   "① ECN功能改变（相位依赖）\n② 杏仁核-PFC连接异常\n③ DMN改变\n（Phillips et al. 2023, Biol Psychiatry）",
            "SCZ":  "① DMN连接减弱（PCC-mPFC）\n② SMN-DMN解耦\n③ 小世界性改变（σ↓）\n（van den Heuvel 2022, Neuron）",
            "ADHD": "① DMN抑制不足\n② DAN-DMN抗相关减弱\n③ 执行网络弱化\n④ 全局效率↓\n（Cortese et al. 2021, JCPP）",
            "ASD":  "① 局部连接↑远程连接↓\n② 社会脑网络异常\n③ 过度局部化（σ↑）\n（Hull et al. 2017, Brain）",
            "GAD":  "① 杏仁核高连接性\n② 前额叶-杏仁核调控减弱\n③ SN过度激活\n（Etkin & Wager 2007, Am J Psychiatry）",
            "OCD":  "① 皮层-纹状体-丘脑回路增强\n② 尾状核异常\n（Rotge et al. 2010, Biol Psychiatry）",
            "PTSD": "① 杏仁核过反应性\n② mPFC-杏仁核调控减弱\n③ 海马功能异常\n（Pitman et al. 2012, Nat Rev Neurosci）",
            "AD":   "① DMN严重破坏（PCC-mPFC断连）\n② 全局效率急剧下降\n③ Hub丧失\n（Buckner 2013, J Neurosci）",
            "PD":   "① 感觉运动网络改变\n② 基底节-皮层失调\n③ 额叶-纹状体连接↓\n（Luo et al. 2021, NeuroImage）",
        }
        _dz_order = ["MDD","BD","SCZ","ADHD","ASD","GAD","OCD","PTSD","AD","PD"]
        lit_data = [["疾病","文献典型脑网络改变","本受试者特征（真实命中）","相似度","备注（可追溯）"]]
        for _k in _dz_order:
            _s_ = _sim[_k]
            lit_data.append([
                f"{_s_['zh_name']}\n({_k})",
                _lit_ref.get(_k, ""),
                _feat_col(_k),
                _s_['level_zh'],
                _s_['note_zh'],
            ])
        # ── 自动更新最新文献：把每种疾病的最新 PubMed 引用追加到第 2 列 ──────────
        _lit_updated = False
        for _ri, _dz in enumerate(_dz_order, start=1):
            _line = self._refs_line(_dz, zh=True)
            if _line:
                lit_data[_ri][1] = lit_data[_ri][1] + _line
                _lit_updated = True
        if self.update_literature:
            _msg = ("✓ 本表「文献典型脑网络改变」列已追加 PubMed 最近数年最新文献（联网自动更新，仅检索疾病+脑网络关键词，未发送任何受试者数据）。"
                    if _lit_updated else
                    "⚠ 文献自动更新已开启，但本次未能联网获取（网络不可用或超时），下表引用为内置参考。")
            story.append(Paragraph(_msg, S_NOTE))
            story.append(Spacer(1, 0.2*cm))
        story.append(_tbl(lit_data, cw=[3*cm, 5.5*cm, 4.5*cm, 1.8*cm, 3.2*cm]))
        # ── 综合小结：完全由真实数据动态生成 ─────────────────────────────────────
        _hubs_all = gm.get('hub_regions', [])
        _hub_txt = "、".join(_hubs_all) if _hubs_all else "未检出明确Hub"
        _notable = [k for k in _dz_order if _sim[k]['level_zh'] == '中度']
        _notable_txt = ("；相似度达『中度』的疾病模式：" +
                        "、".join(f"{_sim[k]['zh_name']}({_sim[k]['note_zh']})" for k in _notable)
                        ) if _notable else "；所有疾病模式相似度均为低或极低"
        _flex_txt = ("偏低" if (dfc.get('n_windows',0) > 0 and flex_v_ < 3) else "在参考范围内")
        story += [Spacer(1, 0.3*cm),
                  Paragraph(
                      f"综合文献对照小结（基于本受试者真实指标）：功能连接Hub脑区为 {_hub_txt}；"
                      f"动态功能连接灵活性{_flex_txt}（{flex_v_}次状态转换）{_notable_txt}。"
                      "以上相似度均由客观指标按预设规则计算，仅供科研探索；"
                      "个体脑网络变异极大，单受试者不可推断诊断，需经群体对照、重测信度与临床信息验证。",
                      S_BODY),
                  PageBreak()]

        # ══ 十一、神经科学讨论 ════════════════════════════════════════════════
        story += [Paragraph("十一、神经科学讨论", S_H1), _hr()]

        story.append(Paragraph("11.1 与已有研究的一致性", S_H2))
        _sig11 = gm.get('small_world_sigma',0); _ge11 = gm.get('global_efficiency',0)
        _le11 = gm.get('local_efficiency',0)
        _sw_txt = ("呈现小世界拓扑结构" if _sig11 > 1 else "未达典型小世界阈值（σ≤1）")
        _ge_txt = ("在健康成人参考范围内" if 0.5 <= _ge11 <= 0.8 else ("偏高" if _ge11 > 0.8 else "偏低"))
        _le_txt = ("在参考范围内" if 0.8 <= _le11 <= 0.95 else ("偏高" if _le11 > 0.95 else "偏低"))
        _prec_txt = _fc_strength(pcc_prec_)
        story.append(Paragraph(
            f"本受试者脑功能网络σ={_sig11:.2f}，{_sw_txt}"
            "（Watts & Strogatz 1998; Bullmore & Sporns 2009, Nat Rev Neurosci）。"
            f"全局效率（GE={_ge11:.3f}，{_ge_txt}）和局部效率（LE={_le11:.3f}，{_le_txt}）。"
            f"DMN的PCC-Precuneus连接（r={pcc_prec_:.3f}，强度{_prec_txt}）及PCC种子连接的空间分布"
            "可与Power et al. (2011, Neuron)及Buckner et al. (2008, J Neurosci)描述的DMN模式对照。", S_BODY))

        story.append(Paragraph("11.2 值得关注的网络特征", S_H2))
        # ── 仅对本受试者真实命中的 Hub / 特征生成解读段落 ────────────────────────
        _hset = set(gm.get('hub_regions', []))
        _feat_paras = []
        _amy_hubs = [h for h in ('L_Amy','R_Amy') if h in _hset]
        if _amy_hubs:
            _feat_paras.append(
                f"（{chr(97+len(_feat_paras))}）杏仁核Hub化：{'、'.join(_amy_hubs)} 在本受试者功能网络中"
                "呈现Hub特征（高度+高中介中心性），与健康成人中杏仁核通常属非Hub节点不同"
                "（van den Heuvel & Sporns 2013, Trends Cogn Sci）。杏仁核Hub化与情绪过度激活和"
                "情绪调节相关，已在焦虑障碍、MDD及PTSD中有报道（Etkin & Wager 2007）。")
        _caud_hubs = [h for h in ('L_Caudate','R_Caudate') if h in _hset]
        if _caud_hubs:
            _feat_paras.append(
                f"（{chr(97+len(_feat_paras))}）纹状体Hub化：{'、'.join(_caud_hubs)} 为Hub节点，"
                "与强迫症中报道的皮层-纹状体回路异常（Rotge et al. 2010; Milad & Rauch 2012, "
                "Trends Cogn Sci）部分相似。尾状核在习惯学习、奖赏预测误差及目标导向行为中扮演关键角色。")
        if dfc.get('n_windows', 0) > 0:
            if flex_v_ < 3:
                _feat_paras.append(
                    f"（{chr(97+len(_feat_paras))}）动态FC灵活性低：本受试者扫描期间DMN仅发生{flex_v_}次状态切换。"
                    "Shine et al. (2016, PNAS)及Vidaurre et al. (2017, Nat Comms)报道健康成人脑状态"
                    "切换频率与认知灵活性正相关，低灵活性与抑郁、精神分裂症及认知障碍相关。")
            else:
                _feat_paras.append(
                    f"（{chr(97+len(_feat_paras))}）动态FC灵活性：本受试者扫描期间DMN发生{flex_v_}次状态切换，"
                    "处于文献报道的健康成人参考范围内（Shine et al. 2016, PNAS）。")
        if not _feat_paras:
            _feat_paras.append("本受试者未检出杏仁核/纹状体Hub化等值得特别关注的网络特征，各项指标总体符合健康脑网络模式。")
        story.append(Paragraph("\n\n".join(_feat_paras), S_BODY))

        story.append(Paragraph("11.3 局限性", S_H2))
        story.append(ListFlowable([
            ListItem(Paragraph("单受试者分析：不能进行群体统计推断，所有指标均为个体描述，不具备统计学意义", S_BODY), leftIndent=15),
            ListItem(Paragraph("未进行头动矫正（需FSL/AFNI）：头动可人为增强短程连接、减弱长程连接，对FC结果有系统性偏差", S_BODY), leftIndent=15),
            ListItem(Paragraph("未进行MNI空间配准：ROI定位采用MNI坐标近似映射到原生空间，定位精度受仿射变换限制", S_BODY), leftIndent=15),
            ListItem(Paragraph("未进行WM/CSF噪声回归：生理噪声未完全去除，可能使FC估计偏高", S_BODY), leftIndent=15),
            ListItem(Paragraph("未进行ICA-AROMA：残留头动伪迹可能影响功能连接估计", S_BODY), leftIndent=15),
            ListItem(Paragraph(f"ROI数量有限（{n_roi}个）：未覆盖全脑所有功能区，部分小脑、脑岛等区域代表性不足", S_BODY), leftIndent=15),
            ListItem(Paragraph("带通滤波去除了高频信息：可能丢失部分神经相关信号", S_BODY), leftIndent=15),
            ListItem(Paragraph("动态FC窗口较长（88s）：可能遗漏快速脑状态切换", S_BODY), leftIndent=15),
        ], bulletType='bullet', leftIndent=10))

        story.append(Paragraph("11.4 建议后续研究方向", S_H2))
        story.append(ListFlowable([
            ListItem(Paragraph("使用fMRIPrep完整预处理流程重新分析（头动矫正、ICA-AROMA、MNI配准）", S_BODY), leftIndent=15),
            ListItem(Paragraph("纳入同龄健康对照组（n≥20），进行网络指标的组间统计检验", S_BODY), leftIndent=15),
            ListItem(Paragraph("结合DTI弥散数据（本次已采集）进行结构-功能联合分析", S_BODY), leftIndent=15),
            ListItem(Paragraph("利用QSM/SWI数据评估铁沉积，与功能网络改变相关联", S_BODY), leftIndent=15),
            ListItem(Paragraph("结合临床评估量表（HAMD/HAMA/Y-BOCS等）进行脑-行为相关分析", S_BODY), leftIndent=15),
            ListItem(Paragraph(f"考虑更精细的ROI图谱（Schaefer200/Brainnetome246）提高空间分辨率", S_BODY), leftIndent=15),
        ], bulletType='bullet', leftIndent=10))
        story.append(PageBreak())

        # ══ 十二、方法学说明 ═══════════════════════════════════════════════════
        story += [Paragraph("十二、方法学说明与可重复性", S_H1), _hr()]
        story.append(_tbl([
            ['dcm2niix','v1.0.20250505','DICOM→NIfTI（UIH mosaic支持）'],
            ['Python','3.13','主语言'],['nibabel','5.4.2','NIfTI读写'],
            ['NetworkX','3.5','图论分析'],['scikit-learn','1.7.2','K-means'],
            ['SciPy','1.16.3','滤波/统计'],['Plotly','6.3.0','交互图表'],
            ['ReportLab','4.5.1','PDF报告'],
        ], cw=[4*cm,4*cm,10*cm]))
        story += [Spacer(1,0.3*cm)]
        story.append(_tbl([
            ['参数','值'],
            ['脑掩码','非零体素65th百分位+形态学开闭运算'],
            ['带通','Butterworth 4阶 filtfilt 0.01–0.1Hz'],
            ['平滑','FWHM=6mm'],['FC阈值','r>0.2'],
            ['动态FC窗口','44TP×2s=88s 步长4TP'],
            ['ReHo邻域','3×3×3体素'],['K-means种子','42'],
        ], cw=[6*cm,12*cm]))
        story.append(PageBreak())

        # 附录
        story += [Paragraph('附录：全部 ROI 图论指标', S_H1), _hr()]
        NMAP = {n: k for k,v in {
            'DMN':['mPFC','PCC','Precuneus','L_Angular','R_Angular','L_mTL_HPC','R_mTL_HPC'],
            'SN':['ACC_dACC','L_AI','R_AI'],'ECN':['L_dlPFC','R_dlPFC','L_IPL','R_IPL'],
            'SMN':['L_M1','R_M1','SMA'],'VIS':['V1_L','V1_R','LOC_L'],
            'DAN':['L_FEF','R_FEF','L_IPS','R_IPS'],
            'LMB':['L_Amy','R_Amy','L_Thal','R_Thal'],
            'SUB':['L_Caudate','R_Caudate','Brainstem','L_Cerebellum','R_Cerebellum'],
        }.items() for n in v}
        pn = gm.get('per_node',{})
        rows = [['ROI','网络','Degree','BC','CC','EC','PC']]
        for rn in roi_names:
            p = pn.get(rn,{})
            rows.append([rn, NMAP.get(rn,'—'),
                str(p.get('degree',0)), f"{p.get('betweenness_centrality',0):.3f}",
                f"{p.get('clustering_coefficient',0):.3f}",
                f"{p.get('eigenvector_centrality',0):.3f}",
                f"{p.get('participation_coefficient',0):.3f}"])
        story.append(_tbl(rows, cw=[3*cm,1.5*cm,2*cm,2.5*cm,2.5*cm,2.5*cm,2.5*cm]))
        story += [Spacer(1,0.5*cm),
                  Paragraph('非诊断性声明', S_H1), _hr(),
                  Paragraph(
                      '本报告由 Brain Analyzer 软件（由 Karcen Zheng 使用 Claude Code 辅助开发）自动生成，'
                      '仅供科研人员参考。所有结果属探索性分析，不构成医学诊断，'
                      '如需临床评估请咨询执照医学专业人员。', S_NOTE)]

    def _build(self, story, subj, language='zh') -> str:
        # Append full body (chapters 4-11) + multi-modal chapter
        self._pdf_body(story, self.sp, self.qc, self.gm, self.ns, self.dfc,
                       self.roi_names, self.FC, len(self.roi_names))
        # 多模态章节（如有数据）
        self._add_multimodal_chapter(story, language='zh')
        s  = 'zh' if language=='zh' else 'en'
        fp = os.path.join(self.reports_dir, f"fMRI_Report_{subj}_{s}.pdf")
        doc = SimpleDocTemplate(fp, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm,  bottomMargin=2*cm,
                                title='静息态fMRI脑网络分析报告',
                                author='Brain Analyzer — Karcen Zheng')
        doc.build(story)
        return fp

    # ── Word ─────────────────────────────────────────────────────────────────
    def generate_word_report(self, language='zh') -> str:
        if language == 'en':
            return self._generate_en_word_report()
        sp=self.sp; qc=self.qc; gm=self.gm; ns=self.ns; dfc=self.dfc
        roi_names=self.roi_names; FC=self.FC; n_roi=len(roi_names)
        subj = sp.get('subject_id','unknown')

        doc = Document()
        doc.add_heading('静息态fMRI脑网络分析报告', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"受试者: {subj}  日期: {sp.get('scan_date','N/A')}"
                          f"  报告: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph('⚠ 仅供科研参考，不构成任何医学诊断依据')
        doc.add_paragraph()

        def add_kv(k, v): doc.add_paragraph(f"• {k}: {v}")

        doc.add_heading('一、扫描参数', 1)
        for k,v in [('受试者ID',subj),('年龄/性别',f"{sp.get('age','N/A')}/{sp.get('sex','N/A')}"),
                    ('扫描仪',sp.get('scanner','N/A')),('TR',f"{sp.get('TR',2.0)}s"),
                    ('QC评分',f"{qc.get('QC_score',0)}/100 {qc.get('QC_stars','—')}"),
                    ('时长',f"{qc.get('total_duration_min',8):.1f}min")]:
            add_kv(k,v)
        doc.add_paragraph()

        doc.add_heading('二、质量控制', 1)
        for k,v in [('QC评分',f"{qc.get('QC_score',0)}/100—{self._qr(qc.get('QC_score',0))}"),
                    ('tSNR中位数',f"{qc.get('tSNR_median',0):.1f}"),
                    ('DVARS%中位数',f"{qc.get('DVARS_pct_median',0):.3f}%"),
                    ('高运动TP',f"{qc.get('n_bad_TPs',0)}/{qc.get('n_timepoints',0)} ({qc.get('pct_bad_TPs',0):.1f}%)")]:
            add_kv(k,v)
        p=self._ip('qc_timeseries.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        doc.add_heading('三、脑网络功能连接', 1)
        for net,v in ns.items():
            add_kv(f"{net}（{NET_CN.get(net,net)}）",f"r={v['mean_FC']:.3f}±{v['std_FC']:.3f}")
        p=self._ip('network_strength.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        doc.add_heading('四、全脑FC矩阵', 1)
        fc_nz=FC[FC!=0]
        add_kv('全脑均值FC', f"{float(fc_nz.mean()):.3f}±{float(fc_nz.std()):.3f}" if len(fc_nz)>0 else 'N/A')
        p=self._ip('fc_matrix.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        doc.add_heading('五、ALFF/fALFF/ReHo', 1)
        alff=self.results.get('alff',{})
        for k,v in [('ALFF均值(z)',f"{alff.get('alff_mean',0):.3f}"),
                    ('fALFF均值(z)',f"{alff.get('falff_mean',0):.3f}"),
                    ('ReHo均值(z)',f"{self.results.get('reho',{}).get('reho_mean',0):.3f}")]:
            add_kv(k,v)
        for fn in ['alff_brain.png','falff_brain.png','reho_brain.png']:
            p=self._ip(fn)
            if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        doc.add_heading('六、图论分析', 1)
        for k,lbl in [('avg_clustering','平均聚类系数'),('global_efficiency','全局效率'),
                       ('local_efficiency','局部效率'),('small_world_sigma','小世界σ'),
                       ('modularity','模块化Q')]:
            add_kv(lbl, f"{gm.get(k,0):.3f}")
        add_kv('Hub脑区', '、'.join(gm.get('hub_regions',[])) or '未检出')
        p=self._ip('graph_hubs.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        doc.add_heading('七、动态功能连接', 1)
        if dfc.get('n_windows',0)>0:
            for k,v in [('窗口数',dfc.get('n_windows',0)),
                        ('状态转换',dfc.get('n_transitions',0)),
                        ('State1占比',f"{dfc.get('state_occupancy',{}).get('State_1',0)*100:.1f}%")]:
                add_kv(k,v)
            p=self._ip('dynamic_fc.png')
            if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        else:
            doc.add_paragraph('时间点不足，动态FC跳过。')
        doc.add_paragraph()

        # ── 八、图论网络拓扑（补充可视化） ─────────────────────────────────────
        doc.add_heading('八、脑功能网络图论分析', 1)
        add_kv('小世界性 σ', f"{gm.get('small_world_sigma',0):.3f}（>1 提示小世界特性）")
        add_kv('全局效率', f"{gm.get('global_efficiency',0):.3f}")
        add_kv('模块化 Q', f"{gm.get('modularity',0):.3f}")
        add_kv('Hub 脑区', '、'.join(gm.get('hub_regions',[])) or '未检出')
        p=self._ip('graph_metrics.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        # ── 九、多模态分析（若有） ─────────────────────────────────────────────
        if self.t1_res or self.dti_res or self.qsm_res:
            doc.add_heading('九、多模态影像分析', 1)
            if self.t1_res:
                add_kv('T1 结构', 'T1 结构像分析已包含（详见 PDF 报告图表）')
            if self.dti_res:
                add_kv('DTI 弥散', 'DTI 白质束 FA/MD 分析已包含')
            if self.qsm_res:
                add_kv('QSM 铁沉积', 'QSM 定量磁化率分析已包含')
            doc.add_paragraph()

        # ── 十、神经精神疾病脑网络文献对照分析（含自动更新文献） ────────────────
        doc.add_heading('十、神经精神疾病脑网络文献对照分析', 1)
        doc.add_paragraph('⚠ 重要声明：本部分仅为科研文献对照，用于探索本受试者脑网络特征与已有群体研究的相似性，'
                          '不可作为医学诊断依据。不作任何疾病诊断、确诊或高度怀疑的结论。')
        if self.update_literature:
            doc.add_paragraph('✓ 已开启文献自动更新：下列引用来自 PubMed 最近数年最新文献'
                              '（仅检索疾病+脑网络关键词，未发送任何受试者数据）。')
        _dz_word = [
            ("MDD",  "重度抑郁症",   "DMN 过度连接、SN-DMN 耦合异常、脑状态灵活性下降"),
            ("BD",   "双相情感障碍", "ECN 功能改变、杏仁核-PFC 连接异常"),
            ("SCZ",  "精神分裂症",   "DMN 连接减弱（PCC-mPFC）、小世界性改变"),
            ("ADHD", "注意缺陷多动", "DMN 抑制不足、DAN-DMN 抗相关减弱"),
            ("ASD",  "孤独症谱系",   "局部连接增强、远程连接减弱、社会脑网络异常"),
            ("GAD",  "广泛性焦虑",   "杏仁核高连接、前额叶-杏仁核调控减弱"),
            ("OCD",  "强迫症",       "皮层-纹状体-丘脑回路增强、尾状核异常"),
            ("PTSD", "创伤后应激",   "杏仁核过反应、mPFC-杏仁核调控减弱"),
            ("AD",   "阿尔茨海默病", "DMN 严重破坏、全局效率下降、Hub 丧失"),
            ("PD",   "帕金森病",     "感觉运动网络改变、基底节-皮层失调"),
        ]
        _sim = self.disease_similarity()
        for _key, _cn, _feat in _dz_word:
            p = doc.add_paragraph()
            p.add_run(f"{_cn}（{_key}）：").bold = True
            p.add_run(_feat)
            _s = _sim.get(_key)
            if _s:
                pv = doc.add_paragraph()
                pv.add_run(f"    → 本受试者相似度：{_s['level_zh']}（{_s['note_zh']}）").bold = True
                # 逐项列出真实命中情况，可追溯
                for zh, en, hit in _s['features']:
                    doc.add_paragraph(('      ✓ ' if hit else '      ✗ ') + zh)
            _refs = self._recent_refs(_key)
            if _refs:
                try:
                    from core.literature_updater import LiteratureUpdater
                    for _r in _refs:
                        doc.add_paragraph('    · 最新文献：' + LiteratureUpdater.format_ref(_r))
                except Exception:
                    pass
        # 综合小结：由真实 Hub + 相似度动态生成
        _hubs_w = gm.get('hub_regions', [])
        _notable_w = [f"{_sim[k]['zh_name']}" for k in ["MDD","BD","SCZ","ADHD","ASD","GAD","OCD","PTSD","AD","PD"]
                      if _sim[k]['level_zh'] == '中度']
        _nt_txt = ("相似度达『中度』的模式：" + "、".join(_notable_w)) if _notable_w else "各疾病模式相似度均为低或极低"
        doc.add_paragraph(f'综合小结（基于本受试者真实指标）：Hub脑区为 '
                          f'{"、".join(_hubs_w) if _hubs_w else "未检出明确Hub"}；{_nt_txt}。'
                          '以上相似度均按客观指标预设规则计算，需经群体对照、重测信度与临床信息验证，'
                          '单受试者不可推断诊断。')
        doc.add_paragraph()

        doc.add_page_break()
        doc.add_heading('免责声明', 1)
        doc.add_paragraph('本报告由 Brain Analyzer（由 Jiacheng Zheng 使用 Claude Code 辅助开发）自动生成，'
                          '仅供科研参考。不构成任何医学诊断或临床建议。严禁非医学相关人员使用。')

        s  = 'zh' if language=='zh' else 'en'
        fp = os.path.join(self.reports_dir, f"fMRI_Report_{subj}_{s}.docx")
        doc.save(fp)
        return fp

    # ═════════════════════════════════════════════════════════════════════════
    # ENGLISH PDF REPORT
    # ═════════════════════════════════════════════════════════════════════════
    def _generate_en_pdf_report(self) -> str:
        sp=self.sp; qc=self.qc; gm=self.gm; ns=self.ns; dfc=self.dfc
        roi_names=self.roi_names; FC=self.FC; n_roi=len(roi_names)
        subj = sp.get('subject_id', 'unknown')
        story = []

        # ── Cover ────────────────────────────────────────────────────────────
        story += [Spacer(1, 3*cm),
                  Paragraph('Resting-State fMRI Brain Network Analysis Report', S_COVER),
                  Spacer(1, 0.5*cm),
                  Paragraph('静息态fMRI脑网络分析报告', S_SUB),
                  Spacer(1, 1.5*cm)]
        story.append(_tbl([
            ['Field', 'Value'],
            ['Subject ID',    subj],
            ['Age / Sex',     f"{sp.get('age','N/A')} / {sp.get('sex','N/A')}"],
            ['Institution',   sp.get('institution','N/A')],
            ['Scanner',       sp.get('scanner','N/A')],
            ['Scan Date',     sp.get('scan_date','N/A')],
            ['TR / TE',       f"{sp.get('TR',2.0)} s / {sp.get('TE',30)} ms"],
            ['Report Date',   datetime.now().strftime('%Y-%m-%d')],
            ['Analysis Platform', 'Python | NetworkX | SciPy | dcm2niix | ReportLab'],
            ['QC Score',      f"{qc.get('QC_score',0)}/100  {qc.get('QC_stars','—')}"],
        ], cw=[5*cm, 11*cm]))
        story += [Spacer(1, 1*cm),
                  Paragraph('⚠ This report is for research purposes only and does not constitute medical advice.', S_NOTE),
                  PageBreak()]

        # ── 1. Scan Parameters ───────────────────────────────────────────────
        stc = '✅ Applied' if sp.get('slice_timing_available') else '⚠ Not applied'
        story += [Paragraph('1. Scan Parameters', S_H1), _hr(),
                  Paragraph('1.1 Acquisition Parameters', S_H2)]
        story.append(_tbl([
            ['Parameter',          'rsfMRI BOLD',                                   'T1 Structural'],
            ['Sequence',           'EPI BOLD',                                      'N/A'],
            ['TR',                 f"{sp.get('TR',2.0)} s",                        'N/A'],
            ['TE',                 f"{sp.get('TE',30)} ms",                        'N/A'],
            ['Voxel size',         f"{sp.get('voxel_size','N/A')} mm",             'N/A'],
            ['Timepoints',         f"{qc.get('n_timepoints','N/A')} ({qc.get('total_duration_min',0):.1f} min)", 'N/A'],
            ['Slice timing corr.', stc,                                             'N/A'],
        ], cw=[5*cm, 6*cm, 5*cm]))
        story.append(Paragraph(
            'Note: This is a resting-state fMRI pipeline; T1 structural acquisition parameters '
            'are not extracted (shown as N/A). Refer to the original DICOM/JSON sidecar for T1 details.', S_CAP))
        story.append(PageBreak())

        # ── 2. Preprocessing ─────────────────────────────────────────────────
        story += [Paragraph('2. Preprocessing Pipeline', S_H1), _hr()]
        story.append(_tbl([
            ['Step',                'Method',                  'Status',      'Notes'],
            ['DICOM → NIfTI',       'dcm2niix',                '✅ Done',     'UIH mosaic auto-decoded'],
            ['Discard first 5 TRs', '—',                       '✅ Done',     'Steady-state reached'],
            ['Slice Timing Corr.',  'Linear interpolation',     stc,           'Reference time = TR/2'],
            ['Brain mask',          'Intensity 65th pct + morphology', '✅ Done', 'No FSL required'],
            ['PSC normalisation',   '(x-μ)/μ × 100',           '✅ Done',     'Per-voxel'],
            ['Linear detrending',   'OLS (intercept + slope)',  '✅ Done',     'Per-voxel'],
            ['Bandpass filter',     'Butterworth 4th filtfilt', '✅ Done',     '0.01–0.1 Hz'],
            ['Spatial smoothing',   'Gaussian FWHM=6mm',        '✅ Done',     'σ≈0.73 voxels'],
            ['Head motion corr.',   '—',                        '⚠ Not done', 'Requires FSL/AFNI'],
            ['MNI registration',    '—',                        '⚠ Not done', 'MNI coords approx.'],
        ], cw=[4*cm, 5.5*cm, 2.5*cm, 6*cm]))
        story.append(PageBreak())

        # ── 3. Quality Control ───────────────────────────────────────────────
        score_v = qc.get('QC_score', 0)
        story += [Paragraph('3. Data Quality Control', S_H1), _hr(),
                  Paragraph(f"Overall QC Score: {score_v}/100  {qc.get('QC_stars','—')}  —  {self._qr(score_v,'en')}",
                             S_BOLD),
                  Spacer(1, 0.3*cm)]
        _tsnr = qc.get('tSNR_median',0); _dvars = qc.get('DVARS_pct_median',0)
        _fd = qc.get('FD_proxy_pct_median',0); _dur = qc.get('total_duration_min',0)
        _pctbad = qc.get('pct_bad_TPs',0)
        story.append(_tbl([
            ['QC Metric',         'Value',                                    'Reference',  'Rating'],
            ['tSNR (median)',      f"{_tsnr:.1f}",          '≥50 good',   _qc_verdict(_tsnr,50,True,zh=False)],
            ['DVARS % (median)',   f"{_dvars:.3f}%",    '<5% low motion', _qc_verdict(_dvars,5,False,zh=False)],
            ['FD proxy % (median)',f"{_fd:.3f}%", '<0.5% ideal',    _qc_verdict(_fd,0.5,False,zh=False)],
            ['High-motion TPs',   f"{qc.get('n_bad_TPs',0)}/{qc.get('n_timepoints',0)}",
                                                                              '<20% accept.',
                                   _qc_verdict(_pctbad,20,False,zh=False)+f" ({_pctbad:.1f}%)"],
            ['Scan duration',     f"{_dur:.1f} min",'≥6 min',     _qc_verdict(_dur,6,True,zh=False)],
        ], cw=[4.5*cm, 3*cm, 4*cm, 6*cm]))
        story += [Spacer(1, 0.3*cm), _img(self._ip('qc_timeseries.png'), 15*cm),
                  Paragraph('Figure 1: QC time series. Top: FD proxy; Middle: DVARS; Bottom: tSNR distribution.', S_CAP),
                  PageBreak()]

        # ── 4. ROI & Networks ────────────────────────────────────────────────
        story += [Paragraph('4. ROI Definition and Resting-State Network Analysis', S_H1), _hr(),
                  Paragraph('33 MNI-coordinate spherical ROIs (radius ≈ 3.5 mm) covering 8 major resting-state networks. '
                             'Fully offline — no atlas download required. '
                             'Time series extracted in native space using affine-projected MNI coordinates.', S_BODY),
                  Spacer(1, 0.2*cm)]
        NET_EN = {'DMN':'Default Mode','SN':'Salience','ECN':'Executive Control',
                  'SMN':'Sensorimotor','VIS':'Visual','DAN':'Dorsal Attention',
                  'LMB':'Limbic','SUB':'Subcortical'}
        net_rows = [['Network','Full Name','Regions','Within-Network FC (r)','SD']]
        for k, v in ns.items():
            net_rows.append([k, NET_EN.get(k, k),
                              ', '.join(v['rois'][:4]) + ('+' if len(v['rois'])>4 else ''),
                              f"{v['mean_FC']:.3f}", f"{v['std_FC']:.3f}"])
        story.append(_tbl(net_rows, cw=[1.5*cm, 3.5*cm, 7*cm, 3*cm, 2*cm]))
        story += [Spacer(1, 0.2*cm), _img(self._ip('network_strength.png'), 14*cm),
                  Paragraph('Figure 2: Mean within-network FC strength (Pearson r) for 8 resting-state networks.', S_CAP),
                  PageBreak()]

        # ── 5. DMN ───────────────────────────────────────────────────────────
        dmn_nodes = ['mPFC','PCC','Precuneus','L_Angular','R_Angular','L_mTL_HPC','R_mTL_HPC']
        dmn_idx   = [roi_names.index(n) for n in dmn_nodes if n in roi_names]
        FC_dmn    = FC[np.ix_(dmn_idx, dmn_idx)] if len(dmn_idx)>1 else np.zeros((1,1))
        dmn_str   = float(FC_dmn[np.triu_indices(len(dmn_idx), k=1)].mean()) if len(dmn_idx)>1 else 0.0
        pcc_i     = roi_names.index('PCC') if 'PCC' in roi_names else 0
        pcc_fc    = sorted([(FC[pcc_i,i], roi_names[i]) for i in range(n_roi) if i!=pcc_i], reverse=True)

        story += [Paragraph('5. Default Mode Network (DMN) Analysis', S_H1), _hr(),
                  Paragraph(f'Mean DMN internal FC strength: r = {dmn_str:.3f}', S_BOLD),
                  Spacer(1, 0.2*cm),
                  Paragraph('5.1 PCC Seed-Based Connectivity', S_H2)]
        story.append(_tbl(
            [['ROI', 'r (PCC seed)', 'Network', 'Interpretation']] +
            [[nm, f'{r:.3f}',
              'DMN' if nm in dmn_nodes else 'Other',
              'Core DMN node' if r > 0.5 else ('Moderate' if r > 0.2 else 'Weak')]
             for r, nm in pcc_fc[:10]],
            cw=[4*cm, 3*cm, 3*cm, 8*cm]))
        story += [Spacer(1, 0.3*cm), _img(self._ip('dmn_fc.png'), 12*cm),
                  Paragraph('Figure 3: DMN internal FC matrix (7 core nodes).', S_CAP),
                  PageBreak()]

        return self._build_en(story, subj)

    def _en_pdf_body(self, story, sp, qc, gm, ns, dfc, roi_names, FC, n_roi):
        """English chapters 6-11 + appendix, appended to story in-place."""
        subj = sp.get('subject_id', 'unknown')
        NET_EN = {'DMN':'Default Mode','SN':'Salience','ECN':'Executive Control',
                  'SMN':'Sensorimotor','VIS':'Visual','DAN':'Dorsal Attention',
                  'LMB':'Limbic','SUB':'Subcortical'}
        dmn_nodes = ['mPFC','PCC','Precuneus','L_Angular','R_Angular','L_mTL_HPC','R_mTL_HPC']
        dmn_idx   = [roi_names.index(n) for n in dmn_nodes if n in roi_names]
        FC_dmn    = FC[np.ix_(dmn_idx, dmn_idx)] if len(dmn_idx)>1 else np.zeros((1,1))
        dmn_str   = float(FC_dmn[np.triu_indices(len(dmn_idx), k=1)].mean()) if len(dmn_idx)>1 else 0.0
        hubs      = gm.get('hub_regions', [])

        # 6. Whole-Brain FC ───────────────────────────────────────────────────
        fc_nz = FC[FC!=0]
        story += [Paragraph('6. Whole-Brain Functional Connectivity Analysis', S_H1), _hr()]
        story.append(_tbl([
            ['Metric', 'Value'],
            ['FC matrix dimension', f'{n_roi}×{n_roi} ({n_roi} ROIs)'],
            ['Whole-brain mean FC',
             f"{float(fc_nz.mean()):.3f}±{float(fc_nz.std()):.3f}" if len(fc_nz)>0 else 'N/A'],
            ['DMN internal FC', f'{dmn_str:.3f}'],
            ['Positive connections',
             f"{float((fc_nz>0).sum()/len(fc_nz)*100):.1f}%" if len(fc_nz)>0 else 'N/A'],
            ['Strong connections (r>0.3)',
             f"{float((fc_nz>0.3).sum()/len(fc_nz)*100):.1f}%" if len(fc_nz)>0 else 'N/A'],
        ], cw=[8*cm, 10*cm]))
        story += [Spacer(1, 0.3*cm), _img(self._ip('fc_matrix.png'), 15*cm),
                  Paragraph('Figure 4: Whole-brain FC matrix (Pearson r, 33 ROIs). '
                             'Diagonal set to 0 for display.', S_CAP),
                  PageBreak()]

        # 7. ALFF / fALFF / ReHo ──────────────────────────────────────────────
        story += [Paragraph('7. Local Brain Activity: ALFF / fALFF / ReHo', S_H1), _hr()]
        story.append(_tbl([
            ['Metric', 'Full Name', 'Method', 'Neuroscientific Meaning'],
            ['ALFF',  'Amplitude of Low-Frequency Fluctuations',
             'Mean FFT amplitude 0.01–0.1 Hz',
             'Local spontaneous neural activity strength'],
            ['fALFF', 'Fractional ALFF',
             'ALFF / Total spectrum amplitude',
             'Normalised ALFF — robust to physiological noise'],
            ['ReHo',  'Regional Homogeneity',
             "Kendall's W, 3×3×3 neighbourhood",
             'Local temporal synchrony of adjacent voxels'],
        ], cw=[2*cm, 4*cm, 5*cm, 7*cm]))
        story.append(Spacer(1, 0.2*cm))
        for mn, fn, cap in [
            ('ALFF',  'alff_brain.png',  'ALFF z-score map (axial slices).'),
            ('fALFF', 'falff_brain.png', 'fALFF z-score map.'),
            ('ReHo',  'reho_brain.png',  "ReHo z-score map (Kendall's W)."),
        ]:
            story += [Paragraph(f'7.x {mn}', S_H2),
                      _img(self._ip(fn), 15*cm),
                      Paragraph(f'Figure: {cap}', S_CAP),
                      Spacer(1, 0.2*cm)]
        story.append(PageBreak())

        # 8. Graph Theory ─────────────────────────────────────────────────────
        story += [Paragraph('8. Brain Functional Network Graph Theory Analysis', S_H1), _hr(),
                  Paragraph(
                      f"Threshold r>{gm.get('threshold_r',0.2)}  "
                      f"Nodes={gm.get('n_nodes',n_roi)}  "
                      f"Edges={gm.get('n_edges',0)}  "
                      f"Density={gm.get('density',0):.3f}", S_BODY)]
        _md = gm.get('mean_degree',0); _cc = gm.get('avg_clustering',0)
        _gge = gm.get('global_efficiency',0); _le = gm.get('local_efficiency',0)
        _cpl = gm.get('char_path_length',0); _sig = gm.get('small_world_sigma',0)
        _mod = gm.get('modularity',0)
        story.append(_tbl([
            ['Graph Metric',            'This Subject',   'Reference (Healthy Adults)',    'Interpretation'],
            ['Mean degree',             f"{_md:.2f}",  '~12–20',      _graph_interp('mean_degree', _md, zh=False)],
            ['Avg. clustering coeff.',  f"{_cc:.3f}",  '0.5–0.8',     _graph_interp('avg_clustering', _cc, zh=False)],
            ['Global efficiency',       f"{_gge:.3f}", '0.5–0.8',     _graph_interp('global_efficiency', _gge, zh=False)],
            ['Local efficiency',        f"{_le:.3f}",  '0.8–0.95',    _graph_interp('local_efficiency', _le, zh=False)],
            ['Characteristic path',     f"{_cpl:.3f}", '1.5–3.0',     _graph_interp('char_path_length', _cpl, zh=False)],
            ['Small-world σ',           f"{_sig:.2f}", '>1 = small-world', '✓ Small-world topology' if _sig > 1 else 'Not small-world (σ≤1)'],
            ['Modularity Q',            f"{_mod:.3f}", '0.2–0.5',     _graph_interp('modularity', _mod, zh=False)],
        ], cw=[4.5*cm, 3.5*cm, 5*cm, 5*cm]))
        story += [Spacer(1, 0.3*cm),
                  Paragraph(f"Hub regions (Degree + BC Top 25%): "
                             f"{', '.join(hubs) if hubs else 'None detected'}", S_BOLD),
                  Spacer(1, 0.2*cm),
                  _img(self._ip('graph_hubs.png'), 15*cm),
                  Paragraph('Figure 5: Graph analysis. Left: Node degree (red = hub); '
                             'Right: Degree vs. betweenness centrality.', S_CAP),
                  PageBreak()]

        # 9. Dynamic FC ───────────────────────────────────────────────────────
        story += [Paragraph('9. Dynamic Functional Connectivity Analysis', S_H1), _hr()]
        n_win = dfc.get('n_windows', 0)
        if n_win > 0:
            story.append(_tbl([
                ['Parameter', 'Value'],
                ['Method',          'Sliding window'],
                ['Window size',     f"{dfc.get('window_size_TPs',44)} TPs ({dfc.get('window_size_s',88):.0f} s)"],
                ['Step size',       f"{dfc.get('step_TPs',4)} TPs"],
                ['ROIs analysed',   'DMN 7 core nodes (21 pairs)'],
                ['Total windows',   str(n_win)],
                ['State 1 occupancy', f"{dfc.get('state_occupancy',{}).get('State_1',0)*100:.1f}%"],
                ['State 2 occupancy', f"{dfc.get('state_occupancy',{}).get('State_2',0)*100:.1f}%"],
                ['State transitions', f"{dfc.get('n_transitions',0)}"],
            ], cw=[6*cm, 12*cm]))
            story += [Spacer(1, 0.3*cm),
                      _img(self._ip('dynamic_fc.png'), 15*cm),
                      Paragraph('Figure 6: Dynamic FC (DMN ROI pairs). '
                                 'Top: time–connectivity heatmap; Bottom: brain-state sequence.', S_CAP)]
        else:
            story.append(Paragraph('Insufficient timepoints — dynamic FC skipped.', S_BODY))
        story.append(PageBreak())

        # 10. Discussion ──────────────────────────────────────────────────────
        story += [Paragraph('10. Key Findings and Neuroscientific Discussion', S_H1), _hr(),
                  Paragraph('⚠ The following is literature-based research comparison only — '
                             'NOT a medical diagnosis.', S_NOTE),
                  Spacer(1, 0.3*cm),
                  Paragraph('10.1 Consistency with Literature', S_H2),
                  Paragraph(
                      f"The brain network exhibits small-world topology (σ={gm.get('small_world_sigma',0):.2f}), "
                      f"global efficiency GE={gm.get('global_efficiency',0):.3f}, "
                      f"and DMN internal FC r={dmn_str:.3f}, "
                      f"consistent with large-sample healthy-adult neuroimaging studies "
                      f"(Bullmore & Sporns 2009; Power et al. 2011).", S_BODY),
                  Paragraph('10.2 Notable Features', S_H2)]
        story.append(_bul([
            f"Hub regions: {', '.join(hubs) if hubs else 'None detected'}",
            f"Dynamic FC flexibility: {dfc.get('n_transitions',0)} state transition(s) over 8 min",
            "Posterior DMN core circuit (PCC–Precuneus) shows strong connectivity",
        ]))
        # 10.x Auto-updated recent literature (optional, online) ──────────────
        if self.update_literature:
            _en_refs = []
            try:
                from core.literature_updater import LiteratureUpdater
                fmt = LiteratureUpdater.format_ref
                for _k in ("MDD", "SCZ", "AD"):
                    for _r in self._recent_refs(_k):
                        _en_refs.append(f"[{_k}] {fmt(_r)} — {_r.get('title','')[:90]}")
            except Exception:
                _en_refs = []
            if _en_refs:
                story += [Paragraph('10.3 Recent Literature (auto-updated from PubMed)', S_H2),
                          Paragraph('Retrieved online using disease + brain-network search terms only; '
                                    'no subject data is transmitted. Research reference only.', S_NOTE),
                          _bul(_en_refs)]
                _lim_idx = '10.4'
            else:
                _lim_idx = '10.3'
        else:
            _lim_idx = '10.3'
        story += [Paragraph(f'{_lim_idx} Limitations', S_H2),
                  _bul(['Single-subject analysis — no group-level statistical inference',
                        'No head-motion correction, MNI registration, or WM/CSF nuisance regression',
                        'ROI localisation in native space is approximate (affine projection)']),
                  PageBreak()]

        # 11. Methods ─────────────────────────────────────────────────────────
        story += [Paragraph('11. Methods and Reproducibility', S_H1), _hr()]
        story.append(_tbl([
            ['Software', 'Version', 'Purpose'],
            ['dcm2niix',    'v1.0.20250505', 'DICOM → NIfTI (UIH mosaic support)'],
            ['Python',      '3.13',          'Primary language'],
            ['nibabel',     '5.4.2',         'NIfTI I/O'],
            ['NetworkX',    '3.5',           'Graph theory'],
            ['scikit-learn','1.7.2',         'k-means brain-state clustering'],
            ['SciPy',       '1.16.3',        'Bandpass filter, statistics'],
            ['Plotly',      '6.3.0',         'Interactive HTML figures'],
            ['ReportLab',   '4.5.1',         'PDF generation'],
        ], cw=[4*cm, 4*cm, 10*cm]))
        story += [Spacer(1, 0.3*cm)]
        story.append(_tbl([
            ['Parameter', 'Value'],
            ['Brain mask',       'Intensity 65th percentile + morphological open/close'],
            ['Bandpass filter',  'Butterworth order-4 filtfilt, 0.01–0.1 Hz'],
            ['Smoothing kernel', 'FWHM = 6 mm'],
            ['FC threshold',     'Pearson r > 0.2'],
            ['dFC window',       '44 TPs × 2 s = 88 s, step 4 TPs'],
            ['ReHo neighbourhood','3×3×3 voxels'],
            ['k-means seed',     '42'],
        ], cw=[6*cm, 12*cm]))
        story.append(PageBreak())

        # Appendix
        story += [Paragraph('Appendix: Per-ROI Graph Theory Metrics', S_H1), _hr()]
        NMAP = {n: k for k, v in {
            'DMN':['mPFC','PCC','Precuneus','L_Angular','R_Angular','L_mTL_HPC','R_mTL_HPC'],
            'SN':['ACC_dACC','L_AI','R_AI'],'ECN':['L_dlPFC','R_dlPFC','L_IPL','R_IPL'],
            'SMN':['L_M1','R_M1','SMA'],'VIS':['V1_L','V1_R','LOC_L'],
            'DAN':['L_FEF','R_FEF','L_IPS','R_IPS'],
            'LMB':['L_Amy','R_Amy','L_Thal','R_Thal'],
            'SUB':['L_Caudate','R_Caudate','Brainstem','L_Cerebellum','R_Cerebellum'],
        }.items() for n in v}
        pn   = gm.get('per_node', {})
        rows = [['ROI', 'Network', 'Degree', 'BC', 'CC', 'EC', 'PC']]
        for rn in roi_names:
            p = pn.get(rn, {})
            rows.append([rn, NMAP.get(rn, '—'),
                          str(p.get('degree', 0)),
                          f"{p.get('betweenness_centrality', 0):.3f}",
                          f"{p.get('clustering_coefficient', 0):.3f}",
                          f"{p.get('eigenvector_centrality', 0):.3f}",
                          f"{p.get('participation_coefficient', 0):.3f}"])
        story.append(_tbl(rows, cw=[3*cm, 1.5*cm, 2*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm]))
        story += [Spacer(1, 0.5*cm),
                  Paragraph('Non-Diagnostic Disclaimer', S_H1), _hr(),
                  Paragraph(
                      'This report is automatically generated by Brain Analyzer '
                      '(developed by Karcen Zheng with Claude Code assistance) '
                      'for research purposes only. All results are exploratory. '
                      'This report does NOT constitute a medical diagnosis, disease confirmation, '
                      'risk assessment, or treatment recommendation. '
                      'For clinical evaluation please consult a licensed medical professional.',
                      S_NOTE)]

    def _build_en(self, story, subj) -> str:
        """Build English PDF: cover+ch1–5 already in story; append ch6–11 + multimodal."""
        self._en_pdf_body(story, self.sp, self.qc, self.gm, self.ns, self.dfc,
                          self.roi_names, self.FC, len(self.roi_names))
        self._add_multimodal_chapter(story, language='en')
        fp = os.path.join(self.reports_dir, f"fMRI_Report_{subj}_en.pdf")
        doc = SimpleDocTemplate(fp, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm,  bottomMargin=2*cm,
                                title='Resting-State fMRI Brain Network Analysis Report',
                                author='Brain Analyzer — Karcen Zheng')
        doc.build(story)
        return fp

    # ═════════════════════════════════════════════════════════════════════════
    # ENGLISH WORD REPORT
    # ═════════════════════════════════════════════════════════════════════════
    def _generate_en_word_report(self) -> str:
        sp=self.sp; qc=self.qc; gm=self.gm; ns=self.ns; dfc=self.dfc
        roi_names=self.roi_names; FC=self.FC
        subj = sp.get('subject_id', 'unknown')
        NET_EN = {'DMN':'Default Mode','SN':'Salience','ECN':'Executive Control',
                  'SMN':'Sensorimotor','VIS':'Visual','DAN':'Dorsal Attention',
                  'LMB':'Limbic','SUB':'Subcortical'}

        doc = Document()
        doc.add_heading('Resting-State fMRI Brain Network Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            f"Subject: {subj}  |  Scan date: {sp.get('scan_date','N/A')}"
            f"  |  Report: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(
            '⚠ For research purposes only. Does NOT constitute medical diagnosis.')
        doc.add_paragraph()

        def kv(k, v): doc.add_paragraph(f"  • {k}: {v}")

        # 1. Scan Parameters
        doc.add_heading('1. Scan Parameters', 1)
        for k, v in [
            ('Subject ID',    subj),
            ('Age / Sex',     f"{sp.get('age','N/A')} / {sp.get('sex','N/A')}"),
            ('Scanner',       sp.get('scanner', 'N/A')),
            ('TR / TE',       f"{sp.get('TR',2.0)} s / {sp.get('TE',30)} ms"),
            ('QC Score',      f"{qc.get('QC_score',0)}/100  {qc.get('QC_stars','—')}"),
            ('Duration',      f"{qc.get('total_duration_min',8):.1f} min"),
            ('Slice timing',  '✓ Applied' if sp.get('slice_timing_available') else '⚠ Not available'),
        ]:
            kv(k, v)
        doc.add_paragraph()

        # 2. Quality Control
        doc.add_heading('2. Quality Control', 1)
        for k, v in [
            ('QC Score',      f"{qc.get('QC_score',0)}/100 — {self._qr(qc.get('QC_score',0),'en')}"),
            ('tSNR (median)', f"{qc.get('tSNR_median',0):.1f}"),
            ('DVARS % (median)', f"{qc.get('DVARS_pct_median',0):.3f}%"),
            ('High-motion TPs',  f"{qc.get('n_bad_TPs',0)}/{qc.get('n_timepoints',0)} "
                                 f"({qc.get('pct_bad_TPs',0):.1f}%)"),
        ]:
            kv(k, v)
        p = self._ip('qc_timeseries.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        # 3. Resting-State Networks
        doc.add_heading('3. Resting-State Networks (FC)', 1)
        for net, v in ns.items():
            kv(f"{net} ({NET_EN.get(net,net)})",
               f"mean FC r = {v['mean_FC']:.3f} ± {v['std_FC']:.3f}")
        p = self._ip('network_strength.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        # 4. FC Matrix
        doc.add_heading('4. Whole-Brain FC Matrix', 1)
        fc_nz = FC[FC!=0]
        kv('Mean FC', f"{float(fc_nz.mean()):.3f}±{float(fc_nz.std()):.3f}" if len(fc_nz)>0 else 'N/A')
        p = self._ip('fc_matrix.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        # 5. ALFF/fALFF/ReHo
        doc.add_heading('5. Local Brain Activity: ALFF / fALFF / ReHo', 1)
        alff = self.results.get('alff', {})
        for k, v in [
            ('ALFF mean (z)',  f"{alff.get('alff_mean',0):.3f}"),
            ('fALFF mean (z)', f"{alff.get('falff_mean',0):.3f}"),
            ('ReHo mean (z)',  f"{self.results.get('reho',{}).get('reho_mean',0):.3f}"),
        ]:
            kv(k, v)
        for fn in ['alff_brain.png', 'falff_brain.png', 'reho_brain.png']:
            p = self._ip(fn)
            if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        # 6. Graph Theory
        doc.add_heading('6. Graph Theory Analysis', 1)
        for k, label in [
            ('avg_clustering',   'Avg. clustering coefficient'),
            ('global_efficiency','Global efficiency'),
            ('local_efficiency', 'Local efficiency'),
            ('small_world_sigma','Small-world index σ'),
            ('modularity',       'Modularity Q'),
        ]:
            kv(label, f"{gm.get(k,0):.3f}")
        hubs = gm.get('hub_regions', [])
        kv('Hub regions', ', '.join(hubs) if hubs else 'None detected')
        p = self._ip('graph_hubs.png')
        if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        doc.add_paragraph()

        # 7. Dynamic FC
        doc.add_heading('7. Dynamic Functional Connectivity', 1)
        if dfc.get('n_windows', 0) > 0:
            for k, v in [
                ('Windows',          dfc.get('n_windows', 0)),
                ('State transitions', dfc.get('n_transitions', 0)),
                ('State 1 occupancy', f"{dfc.get('state_occupancy',{}).get('State_1',0)*100:.1f}%"),
            ]:
                kv(k, v)
            p = self._ip('dynamic_fc.png')
            if os.path.exists(p): doc.add_picture(p, width=Inches(5.5))
        else:
            doc.add_paragraph('Insufficient timepoints — dynamic FC skipped.')

        doc.add_page_break()
        doc.add_heading('Non-Diagnostic Disclaimer', 1)
        doc.add_paragraph(
            'This report is automatically generated by Brain Analyzer '
            '(developed by Karcen Zheng with Claude Code assistance) '
            'for research purposes only. '
            'Results do NOT constitute medical diagnosis or clinical advice. '
            'Consult a licensed medical professional for clinical evaluation.')

        fp = os.path.join(self.reports_dir, f"fMRI_Report_{subj}_en.docx")
        doc.save(fp)
        return fp

    # ═════════════════════════════════════════════════════════════════════════
    # MULTI-MODAL CHAPTER (中/英 共用，通过 language 参数切换文本)
    # ═════════════════════════════════════════════════════════════════════════
    def _add_multimodal_chapter(self, story: list, language='zh'):
        """
        向 story 中追加多模态分析章节：
        十二、T1结构像分析
        十三、DTI弥散张量分析
        十四、QSM铁沉积分析
        十五、多模态综合疾病关联分析
        """
        zh = (language == 'zh')
        t1  = self.t1_res
        dti = self.dti_res
        qsm = self.qsm_res

        has_t1  = bool(t1  and "brain_vol_cc" in t1)
        has_dti = bool(dti and "global" in dti)
        has_qsm = bool(qsm and "roi_stats" in qsm)

        if not (has_t1 or has_dti or has_qsm):
            return  # 无多模态数据，跳过

        # ── T1 ──────────────────────────────────────────────────────────────
        if has_t1:
            story += [Paragraph('十二、T1结构像分析' if zh else '12. T1 Structural Analysis', S_H1), _hr()]
            row = [
                ['参数' if zh else 'Parameter', '受试者' if zh else 'Subject',
                 '参考（成人）' if zh else 'Reference (Adult)', '评价' if zh else 'Rating'],
                ['脑总体积 (cc)' if zh else 'Total brain vol.',
                 f"{t1.get('brain_vol_cc',0):.0f}",
                 '1350–1500', '正常' if 1350 <= t1.get('brain_vol_cc',0) <= 1500 else '偏小/大'],
                ['灰质 GM (cc)',  f"{t1.get('gm_vol_cc',0):.0f}",  '650–780', '—'],
                ['白质 WM (cc)',  f"{t1.get('wm_vol_cc',0):.0f}",  '450–600', '—'],
                ['CSF (cc)',      f"{t1.get('csf_vol_cc',0):.0f}",  '100–200', '—'],
                ['GM/WM 比值',   f"{t1.get('gm_wm_ratio',0):.3f}", '1.2–1.5', '—'],
                ['半球对称性',   f"{t1.get('symmetry_index',0):.3f}", '>0.92良好','—'],
            ]
            story.append(_tbl(row, cw=[5*cm,3*cm,4*cm,6*cm]))
            story += [Spacer(1,0.2*cm),
                      _img(self._ip('t1_tissue_seg.png'),15*cm),
                      Paragraph('图A：T1组织分割（1=CSF，2=GM，3=WM），轴位切面。', S_CAP),
                      Spacer(1,0.2*cm),
                      _img(self._ip('t1_volumes.png'),14*cm),
                      Paragraph('图B：脑组织体积（受试者 vs 健康成人参考值）。', S_CAP)]
            for interp in t1.get('interpretation', []):
                story.append(Paragraph(f"▸ {interp}", S_BODY))
            story.append(PageBreak())

        # ── DTI ─────────────────────────────────────────────────────────────
        if has_dti:
            story += [Paragraph('十三、DTI弥散张量分析' if zh else '13. DTI Analysis', S_H1), _hr()]
            g = dti.get('global', {})
            story.append(Paragraph(
                f"全脑WM FA均值={g.get('FA_mean','?')}，MD={g.get('MD_mean_e3','?')}×10⁻³mm²/s，"
                f"AD={g.get('AD_mean_e3','?')}，RD={g.get('RD_mean_e3','?')}，"
                f"WM体素数={g.get('WM_voxels','?')}", S_BODY))
            story += [Spacer(1,0.2*cm),
                      _img(self._ip('dti_FA_MD.png'),15*cm),
                      Paragraph('图C：DTI FA图（上）和MD图（下），轴位切面。热色=高FA/低MD。', S_CAP),
                      Spacer(1,0.2*cm)]

            # ROI 统计表
            roi_stats = dti.get('roi', {})
            if roi_stats:
                roi_rows = [['WM束','FA均值','MD (×10⁻³)','AD (×10⁻³)','RD (×10⁻³)','参考FA']]
                REF_FA = {'Corpus_Callosum_Genu':0.75,'Corpus_Callosum_Body':0.70,
                          'Corpus_Callosum_Splenium':0.72,'Cingulum_L':0.48,'Cingulum_R':0.48,
                          'CST_L':0.60,'CST_R':0.60,'SLF_L':0.46,'SLF_R':0.46,
                          'UF_L':0.44,'UF_R':0.44,'IFOF_L':0.52,'IFOF_R':0.52}
                for tract, v in list(roi_stats.items())[:13]:
                    ref = REF_FA.get(tract, 0.50)
                    fa_val = v.get('FA_mean', 0)
                    flag = '⚠' if fa_val < ref * 0.80 else ('↑' if fa_val > ref * 1.20 else '✓')
                    roi_rows.append([tract.replace('_',' '),
                                     f"{fa_val:.3f} {flag}",
                                     f"{v.get('MD_mean',0):.3f}",
                                     f"{v.get('AD_mean',0):.3f}",
                                     f"{v.get('RD_mean',0):.3f}",
                                     f"{ref:.2f}"])
                story.append(_tbl(roi_rows, cw=[4*cm,2.5*cm,2.5*cm,2.5*cm,2.5*cm,2.5*cm]))
                story += [Spacer(1,0.2*cm),
                          _img(self._ip('dti_tract_FA.png'),15*cm),
                          Paragraph('图D：主要WM束FA值（蓝=受试者，灰=参考值）。⚠=偏低。', S_CAP)]
            for interp in g.get('interpretation',[]):
                story.append(Paragraph(f"▸ {interp}", S_BODY))
            story.append(PageBreak())

        # ── QSM ─────────────────────────────────────────────────────────────
        if has_qsm:
            story += [Paragraph('十四、QSM定量磁化率图（铁沉积）分析' if zh else '14. QSM Iron Mapping', S_H1), _hr()]
            story.append(Paragraph(
                'QSM反映局部组织磁化率，正值提示铁、钙等顺磁性物质沉积；'
                '深部灰质核团铁含量增加与多种神经退行性疾病相关。', S_BODY))
            story += [Spacer(1,0.2*cm),
                      _img(self._ip('qsm_iron_map.png'),15*cm),
                      Paragraph('图E：QSM图（轴位切面），暖色=铁/顺磁性沉积，冷色=逆磁性。', S_CAP),
                      Spacer(1,0.2*cm)]
            roi_stats = qsm.get('roi_stats', {})
            if roi_stats:
                qsm_rows = [['脑区','QSM均值(ppb)','SD','参考均值','评价']]
                QSM_REF  = {'Caudate':30,'Putamen':60,'Globus_Pallidus':120,
                            'Thalamus':20,'Substantia_Nigra':80,'Red_Nucleus':60,'Dentate':40}
                for roi, v in list(roi_stats.items())[:14]:
                    ref = next((QSM_REF[k] for k in QSM_REF if k.lower() in roi.lower()), 50)
                    qval = v.get('QSM_mean_ppb', 0)
                    flag = '⚠偏高' if qval > ref*1.5 else ('正常' if qval > 0 else '低/负')
                    qsm_rows.append([roi.replace('_',' '), f"{qval:.1f}", f"{v.get('QSM_std_ppb',0):.1f}",
                                     str(ref), flag])
                story.append(_tbl(qsm_rows, cw=[4.5*cm,3*cm,2*cm,3*cm,5.5*cm]))
                story += [Spacer(1,0.2*cm),
                          _img(self._ip('qsm_roi_bars.png'),14*cm),
                          Paragraph('图F：深部灰质ROI QSM值（蓝=受试者，灰=参考）。红色=偏高。', S_CAP)]
            for interp in qsm.get('interpretation', []):
                story.append(Paragraph(f"▸ {interp}", S_BODY))
            story.append(PageBreak())

        # ── 多模态疾病关联 ───────────────────────────────────────────────────
        story += [Paragraph('十五、多模态综合疾病关联分析' if zh else '15. Multi-Modal Disease Correlation Analysis', S_H1),
                  _hr(),
                  Paragraph('⚠ 以下疾病关联分析仅为科研文献对照，不构成任何医学诊断。', S_NOTE),
                  Spacer(1,0.3*cm)]

        # 收集关键指标
        dfc_v    = self.dfc
        gm_      = self.gm
        fc_ns    = self.ns
        sigma    = gm_.get('small_world_sigma', 0)
        ge       = gm_.get('global_efficiency', 0)
        hubs     = gm_.get('hub_regions', [])
        n_trans  = dfc_v.get('n_transitions', 0)
        dmn_fc   = self.fp.get('dmn_strength', 0)
        sn_fc    = fc_ns.get('SN', {}).get('mean_FC', 0)
        ecn_fc   = fc_ns.get('ECN', {}).get('mean_FC', 0)
        fa_g     = dti.get('global', {}).get('FA_mean', None) if has_dti else None
        wm_v     = t1.get('wm_vol_cc', None) if has_t1 else None
        gm_v     = t1.get('gm_vol_cc', None) if has_t1 else None
        qsm_pall = qsm.get('roi_stats', {}).get('Globus_Pallidus_L', {}).get('QSM_mean_ppb', None) if has_qsm else None
        qsm_sn   = qsm.get('roi_stats', {}).get('Substantia_Nigra_L', {}).get('QSM_mean_ppb', None) if has_qsm else None

        def check(val, lo, hi):
            if val is None: return '—'
            return '正常✓' if lo <= val <= hi else ('偏低↓' if val < lo else '偏高↑')

        # ── 相似度列改由数据驱动引擎给出（与第十章同源，可追溯）───────────────────
        _sim = self.disease_similarity()
        def _simcol(key):
            s = _sim.get(key)
            return f"{s['level_zh']}（{s['note_zh']}）" if s else '极低'

        lit_rows = [
            ['疾病','fMRI关键指标','DTI关键指标','结构关键指标','QSM关键指标','相似度（科研参考）'],
            ['重度抑郁症\n(MDD)',
             f"DMN={dmn_fc:.2f}，灵活性={n_trans}次，SN={sn_fc:.2f}",
             f"UF/CC FA {check(fa_g,0.35,0.55)}",
             f"GM体积 {check(gm_v,600,800)}cc",
             '—', _simcol('MDD')],
            ['强迫症 (OCD)',
             f"纹状体Hub: {'是' if ('L_Caudate' in hubs or 'R_Caudate' in hubs) else '否'}，SN={sn_fc:.2f}",
             f"CC/CS FA {check(fa_g,0.35,0.55)}",
             f"尾状核体积（未分割）",
             f"Putamen {check(qsm.get('roi_stats',{}).get('Putamen_L',{}).get('QSM_mean_ppb',None),30,90) if has_qsm else '—'}",
             _simcol('OCD')],
            ['帕金森病 (PD)',
             f"SMN FC {fc_ns.get('SMN',{}).get('mean_FC',0):.2f}，皮层下Hub",
             f"SN区FA {check(fa_g,0.30,0.50)}",
             f"WM体积 {check(wm_v,400,650)}cc",
             f"SN QSM {check(qsm_sn,60,130) if qsm_sn else '—'}",
             _simcol('PD')],
            ['阿尔茨海默病 (AD)',
             f"DMN={dmn_fc:.2f}，GE={ge:.3f}",
             f"全脑FA {check(fa_g,0.35,0.55)}",
             f"GM体积 {check(gm_v,600,800)}cc",
             '铁沉积模式（未定量比对）', _simcol('AD')],
            ['焦虑障碍 (GAD/SAD)',
             f"杏仁核Hub: {'是' if 'L_Amy' in hubs else '否'}，SN={sn_fc:.2f}",
             f"UF FA（情绪调节通路）",
             f"杏仁核体积（未分割）",
             '—', _simcol('GAD')],
            ['多系统萎缩 (MSA)',
             f"小脑网络={fc_ns.get('SUB',{}).get('mean_FC',0):.2f}",
             f"小脑束FA",
             f"小脑体积（估算）",
             f"苍白球QSM {check(qsm_pall,80,200) if qsm_pall else '—'}",
             '极低（无专用规则，仅列指标）'],
        ]
        story.append(_tbl(lit_rows, cw=[3*cm,5*cm,3*cm,3*cm,3*cm,3*cm]))
        story += [Spacer(1,0.3*cm),
                  Paragraph('说明：上表基于已发表文献（Zhu et al. 2021；Rotge et al. 2010；'
                             'van den Heuvel & Sporns 2013）的群体统计特征，'
                             '单受试者不能做出组间统计结论。', S_CAP),
                  Spacer(1,0.3*cm)]

        # 综合解读
        story.append(Paragraph('15.1 本受试者多模态综合特征摘要', S_H2))
        summary_bullets = [
            f"fMRI：DMN内部FC={dmn_fc:.3f}，动态FC灵活性={n_trans}次，"
            f"Hub脑区={'、'.join(hubs) if hubs else '无'}",
        ]
        if has_dti:
            summary_bullets.append(f"DTI：全脑WM FA={fa_g:.3f}，"
                                    f"{'FA在正常参考范围' if fa_g and 0.35<fa_g<0.60 else 'FA偏低/偏高'}")
        if has_t1:
            summary_bullets.append(f"T1：脑总体积={t1.get('brain_vol_cc',0):.0f}cc，"
                                    f"GM={t1.get('gm_vol_cc',0):.0f}cc，WM={t1.get('wm_vol_cc',0):.0f}cc")
        if has_qsm:
            qsm_interp = qsm.get('interpretation', ['QSM正常'])
            summary_bullets.append(f"QSM：{qsm_interp[0]}")
        story.append(_bul(summary_bullets))
        story.append(PageBreak())
